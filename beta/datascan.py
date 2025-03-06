import os
import json
import torch
import tkinter as tk
from tkinter import filedialog, messagebox, simpledialog, ttk
from PyPDF2 import PdfReader
from docx import Document
from openpyxl import load_workbook
from transformers import AutoTokenizer, AutoModelForSeq2SeqLM, pipeline
from bs4 import BeautifulSoup
import validators
import requests
import threading
from queue import Queue
from urllib.parse import urlparse
import time
import pandas as pd
from datetime import datetime
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter
from reportlab.platypus import Table, TableStyle
from reportlab.lib import colors
import matplotlib.pyplot as plt
from matplotlib.backends.backend_tkagg import FigureCanvasTkAgg

class Config:
    MODEL_NAME = "google/flan-t5-small"
    MAX_TOKENS = 384
    MAX_CONTEXT_LENGTH = 3
    USE_QUANTIZATION = True
    DEVICE = "cuda" if torch.cuda.is_available() else "cpu"
    TEMP_DIR = "temp_processing"
    BACKUP_DIR = "backups"
    
     # Nova paleta de cores profissional
    PRIMARY_COLOR = "#2A3F5F"      # Azul escuro
    SECONDARY_COLOR = "#4A6572"    # Cinza azulado
    ACCENT_COLOR = "#F9AA33"       # Laranja dourado
    BACKGROUND_COLOR = "#E0E0E0"   # Cinza claro
    TEXT_COLOR = "#FFFFFF"         # Branco
    CHAT_BG = "#37474F"            # Cinza escuro
    HOVER_COLOR = "#566573"        # Efeito hover
    

class DocumentProcessor:
    def __init__(self):
        self.document_content = ""
        self.scraped_data = {}
        self.chat_history = []
        
        try:
            self.tokenizer = AutoTokenizer.from_pretrained(Config.MODEL_NAME)
            
            self.model = AutoModelForSeq2SeqLM.from_pretrained(
                Config.MODEL_NAME,
                device_map="auto" if torch.cuda.is_available() else None,
                torch_dtype=torch.float32,
                low_cpu_mem_usage=False
            )
            
            if Config.USE_QUANTIZATION and torch.cuda.is_available():
                self.model = self.quantize_model(self.model)
                
            self.model.to(Config.DEVICE)
            self.model.eval()

            self.generator = pipeline(
                'text2text-generation',
                model=self.model,
                tokenizer=self.tokenizer,
                max_length=Config.MAX_TOKENS
            )
            
        except Exception as e:
            raise Exception(f"Falha ao carregar modelo: {str(e)}")

    def quantize_model(self, model):
        return torch.quantization.quantize_dynamic(
            model,
            {torch.nn.Linear},
            dtype=torch.qint8
        )

    def extract_from_file(self, file_path: str) -> str:
        try:
            chunk_size = 1024
            content = []
            ext = os.path.splitext(file_path)[1].lower()

            if ext == '.pdf':
                reader = PdfReader(file_path)
                for page in reader.pages[:10]:
                    text = page.extract_text() or ""
                    for i in range(0, len(text), chunk_size):
                        content.append(text[i:i+chunk_size])

            elif ext == '.docx':
                doc = Document(file_path)
                for para in doc.paragraphs[:100]:
                    if para.text.strip():
                        content.append(para.text[:chunk_size].replace('\n', ' '))

            elif ext in ('.xlsx', '.xls'):
                wb = load_workbook(file_path, read_only=True)
                for sheet in wb:
                    for row in sheet.iter_rows(values_only=True)[:1000]:
                        row_content = [str(cell)[:50] for cell in row if cell is not None]
                        content.append(" | ".join(row_content))
                    break

            elif ext == '.txt':
                with open(file_path, 'r', encoding='utf-8', errors='replace') as f:
                    for _ in range(1000):
                        chunk = f.read(chunk_size)
                        if not chunk:
                            break
                        content.append(chunk.replace('\n', ' '))
            else:
                raise ValueError(f"Formato não suportado: {ext}")

            full_content = " ".join(content)
            return full_content[:Config.MAX_TOKENS * 4].strip()

        except Exception as e:
            raise Exception(f"Erro na extração: {str(e)}")

    def export_to_pdf(self, data: dict, filename: str):
        """Exporta dados para PDF com formatação profissional"""
        try:
            c = canvas.Canvas(filename, pagesize=letter)
            width, height = letter
            
            # Cabeçalho
            c.setFillColorRGB(0, 1, 0)  # Verde
            c.setFont("Helvetica-Bold", 16)
            c.drawString(100, height - 100, "Relatório Gerado por DataScan")
            c.line(100, height - 110, width - 100, height - 110)
            
            # Conteúdo
            y_position = height - 150
            c.setFont("Helvetica", 12)
            c.setFillColorRGB(1, 1, 1)  # Branco
            
            # Tabela de dados
            if 'data' in data and len(data['data']) > 0:
                table_data = [list(data['data'][0].keys())]  # Cabeçalhos
                for row in data['data']:
                    table_data.append([str(v) for v in row.values()])
                
                t = Table(table_data)
                t.setStyle(TableStyle([
                    ('BACKGROUND', (0,0), (-1,0), colors.green),
                    ('TEXTCOLOR', (0,0), (-1,0), colors.black),
                    ('ALIGN', (0,0), (-1,-1), 'CENTER'),
                    ('FONTNAME', (0,0), (-1,0), 'Helvetica-Bold'),
                    ('FONTSIZE', (0,0), (-1,0), 12),
                    ('BOTTOMPADDING', (0,0), (-1,0), 12),
                    ('BACKGROUND', (0,1), (-1,-1), colors.black),
                    ('TEXTCOLOR', (0,1), (-1,-1), colors.green),
                    ('GRID', (0,0), (-1,-1), 1, colors.green)
                ]))
                
                t.wrapOn(c, width-200, height)
                t.drawOn(c, 100, y_position - len(table_data)*20)
                y_position -= len(table_data)*20 + 50

            # Análise da IA
            if 'analysis' in data:
                c.setFont("Helvetica", 10)
                c.drawString(100, y_position, "Análise da IA:")
                text_object = c.beginText(100, y_position - 20)
                text_object.setFont("Helvetica", 10)
                text_object.setFillColor(colors.green)
                
                for line in data['analysis'].split('\n'):
                    text_object.textLine(line)
                    y_position -= 12
                    if y_position < 100:
                        c.showPage()
                        y_position = height - 100
                        text_object = c.beginText(100, y_position)
                        
                c.drawText(text_object)

            c.save()
            return True
        except Exception as e:
            raise Exception(f"Erro ao gerar PDF: {str(e)}")

    def scrape_website(self, url: str) -> dict:
        try:
            url = url.strip().replace(" ", "")
            if not url:
                raise ValueError("URL não pode estar vazia")
            if not url.startswith(('http://', 'https://')):
                url = f'https://{url}'
            if not validators.url(url):
                raise ValueError(f"URL inválida: {url}")
            parsed_url = urlparse(url)
            if not validators.domain(parsed_url.hostname):
                raise ValueError(f"Domínio inválido: {parsed_url.hostname}")
            response = requests.get(url, timeout=15)
            response.raise_for_status()
            soup = BeautifulSoup(response.text, 'html.parser')
            raw_text = soup.get_text(separator=' ', strip=True)[:4000]
            
            # Processamento com IA para extração estruturada
            ai_prompt = (
                "Extraia informações estruturadas deste site. Campos importantes podem incluir: "
                "nomes, valores, datas, descrições. Formato desejado: JSON com chaves e valores. "
                f"Conteúdo:\n{raw_text}\n\nResposta:"
            )
            
            ai_response = self.generate_ai_response(ai_prompt)
            
            # Tenta converter a resposta da IA para dicionário
            try:
                structured_data = json.loads(ai_response)
            except json.JSONDecodeError:
                structured_data = {"content": ai_response}
            
            # Combina com dados tradicionais
            data = {
                "title": soup.title.string if soup.title else "Sem título",
                "headers": [header.text.strip() for header in soup.find_all(['h1', 'h2', 'h3'])],
                "links": [link.get('href') for link in soup.find_all('a') if link.get('href')],
                "content": raw_text,
                "structured_data": structured_data,
                "ai_analysis": ai_response
            }
            
            return data
        except Exception as e:
            raise Exception(f"Erro no scraping: {str(e)}")

    def generate_ai_response(self, prompt: str) -> str:
        try:
            # Tratamento para saudações básicas
            greetings = {
                "olá": "Olá! Como posso ajudá-lo hoje?",
                "oi": "Oi! Em que posso ser útil?",
                "bom dia": "Bom dia! O que gostaria de fazer hoje?",
                "boa tarde": "Boa tarde! Como posso ajudá-lo?",
                "boa noite": "Boa noite! Em que posso ser útil?"
            }

            lower_prompt = prompt.lower()
            if lower_prompt in greetings:
                return greetings[lower_prompt]

            # Construir o contexto da conversa
            context = "\n".join(self.chat_history[-Config.MAX_CONTEXT_LENGTH:])

            full_prompt = (
                f"Diálogo anterior:\n{context}\n\n"
                f"Usuário: {prompt}\n"
                f"Assistente:"
            )

            # Configurações de geração
            generation_config = {
                'max_new_tokens': 200,
                'temperature': 0.7,
                'repetition_penalty': 1.2,
                'num_beams': 2,
                'early_stopping': True,
                'no_repeat_ngram_size': 2,
                'do_sample': True  # Usar temperature para amostragem
            }

            # Preparar o modelo e tokenizer para a entrada
            inputs = self.tokenizer(
                full_prompt,
                return_tensors="pt",
                max_length=Config.MAX_TOKENS,
                truncation=True
            ).to(Config.DEVICE)

            with torch.no_grad():
                outputs = self.model.generate(
                    **inputs,
                    **generation_config
                )

            response = self.tokenizer.decode(
                outputs[0],
                skip_special_tokens=True
            ).strip()

            return response

        except Exception as e:
            return f"Erro na geração: {str(e)}"

    def generate_template(self, template_type: str, fields: list) -> str:
        try:
            # Validar tipo de template
            template_type = template_type.lower()
            if template_type not in ['docx', 'xlsx']:
                raise ValueError("Tipo de template inválido. Use 'docx' ou 'xlsx'.")

            if not fields:
                raise ValueError("A lista de campos não pode estar vazia.")

            # Criar prompt para geração de template
            prompt = (
                f"Crie um template {template_type} com os seguintes campos: {', '.join(fields)}. "
                "Inclua marcações {{campo}} para preenchimento. Formato exemplo:"
                "\n\nPara DOCX:\n[Nome: {{nome}}]\n[Data: {{data}}]"
                "\n\nPara Excel:\n| Nome | Data | Valor |"
            )

            response = self.generate_ai_response(prompt)

            # Definir caminho de saída
            ext = template_type
            output_path = os.path.join(Config.TEMP_DIR, f"template_{int(time.time())}.{ext}")

            # Gerar arquivo conforme o tipo
            if template_type == 'docx':
                doc = Document()
                for line in response.split('\n'):
                    doc.add_paragraph(line)
                doc.save(output_path)

            elif template_type == 'xlsx':
                wb = Workbook()
                ws = wb.active
                headers = [field.strip() for field in fields]
                ws.append(headers)
                wb.save(output_path)

            return output_path

        except ValueError as ve:
            raise ValueError(f"Erro de validação: {str(ve)}")
        except Exception as e:
            raise Exception(f"Erro na geração do template: {str(e)}")

    def fill_template(self, template_path: str, output_path: str, data: dict) -> bool:
        try:
            # Combina dados tradicionais e estruturados
            full_data = {
                **data.get('structured_data', {}),
                **{f"raw_{k}": v for k, v in data.items()}
            }
            
            ext = os.path.splitext(template_path)[1].lower()
            if ext == '.docx':
                self.fill_docx_template(template_path, output_path, full_data)
            elif ext in ('.xlsx', '.xls'):
                self.fill_excel_template(template_path, full_data).save(output_path)
            elif ext == '.pdf':
                raise NotImplementedError("Preenchimento de PDF não implementado")
            return True
        except Exception as e:
            raise Exception(f"Erro no template: {str(e)}")

    def fill_docx_template(self, template_path: str, output_path: str, data: dict):
        doc = Document(template_path)
        for para in doc.paragraphs:
            if '{{' in para.text:
                for key, value in data.items():
                    if isinstance(value, list):
                        value = ', '.join(value)
                    para.text = para.text.replace(f'{{{{{key}}}}}', str(value))
        doc.save(output_path)

    def fill_excel_template(self, template_path: str, data: dict):
        wb = load_workbook(template_path)
        for sheet in wb:
            for row in sheet.iter_rows():
                for cell in row:
                    if cell.value and '{{' in str(cell.value):
                        for key, value in data.items():
                            if isinstance(value, list):
                                value = ', '.join(value)
                            cell.value = str(cell.value).replace(f'{{{{{key}}}}}', str(value))
        return wb


class Application:
    def __init__(self):
        self.root = tk.Tk()
        self.root.title("DataScan")
        self.root.geometry("1200x800")
        self.root.configure(bg=Config.BACKGROUND_COLOR)
        self.status_var = tk.StringVar()
        self.process_file_btn = None
        self.scrape_btn = None
        self.chat_btn = None
        self.template_btn = None
        
        self.processor = DocumentProcessor()
        self.ui_queue = Queue()
        self.progress_window = None
        self.progress_bar = None
        
        self.setup_ui()
        self.setup_ui_handler()
        self.current_chart_window = None
        self.tree = None
        self.setup_db_explorer()

    def setup_ui(self):
        main_frame = tk.Frame(self.root, bg=Config.BACKGROUND_COLOR)
        main_frame.pack(fill=tk.BOTH, expand=True, padx=20, pady=20)

        # Frame centralizado
        center_frame = tk.Frame(main_frame, bg=Config.BACKGROUND_COLOR)
        center_frame.place(relx=0.5, rely=0.5, anchor='center')

        # Título moderno
        title_frame = tk.Frame(center_frame, bg=Config.BACKGROUND_COLOR)
        title_frame.pack(pady=20)
        
        self.title_label = tk.Label(
            title_frame,
            text="DataScan",
            font=("Segoe UI", 34, "bold"),
            fg=Config.PRIMARY_COLOR,
            bg=Config.BACKGROUND_COLOR
        )
        self.title_label.pack()
        
        self.subtext_label = tk.Label(
            title_frame,
            text="Intelligent Data Processing",
            font=("Segoe UI", 18),
            fg=Config.SECONDARY_COLOR,
            bg=Config.BACKGROUND_COLOR
        )
        self.subtext_label.pack()
        
            # Adicionar barra de status
        self.status_bar = tk.Label(
            self.root, 
            textvariable=self.status_var,
            bg=Config.SECONDARY_COLOR,
            fg=Config.TEXT_COLOR,
            anchor=tk.W
        )
        self.status_bar.pack(side=tk.BOTTOM, fill=tk.X)
        
        self.subtext2_label = tk.Label(
            title_frame,
            text="By: github.com/Marlon009",
            font=("Segoe UI", 12, "italic"),
            fg=Config.SECONDARY_COLOR,
            bg=Config.BACKGROUND_COLOR
        )
        self.subtext2_label.pack()

        # Botões principais
        button_frame = tk.Frame(center_frame, bg=Config.BACKGROUND_COLOR)
        button_frame.pack(pady=30)

        button_style = {
            'font': ('Segoe UI', 12),
            'width': 20,
            'height': 2,
            'bg': Config.PRIMARY_COLOR,
            'fg': Config.TEXT_COLOR,
            'activebackground': Config.HOVER_COLOR,
            'activeforeground': Config.TEXT_COLOR,
            'relief': 'flat',
            'borderwidth': 0,
            'cursor': 'hand2'
        }

        buttons = [
            ("📁 Processar Arquivo", self.process_file),
            ("🌐 Scraping Web", self.scrape_website),
            ("💬 Terminal IA", self.chat_with_ai),
            ("💾 Backup", self.create_backup)
        ]

        for text, cmd in buttons:
            btn = tk.Button(button_frame, text=text, compound=tk.LEFT, **button_style)
            btn.config(command=cmd)
            btn.bind("<Enter>", lambda e: e.widget.config(bg=Config.ACCENT_COLOR))
            btn.bind("<Leave>", lambda e: e.widget.config(bg=Config.PRIMARY_COLOR))
            btn.pack(side=tk.LEFT, padx=10, pady=5)
        

    def show_progress(self, message: str):
        if self.progress_window is None:
            self.progress_window = tk.Toplevel(self.root)
            self.progress_window.title("Processando")
            self.progress_window.configure(bg='#0a0a0a')
            
            self.progress_label = tk.Label(
                self.progress_window,
                text=message,
                font=("Segoe UI", 12),
                fg="#00ff00",
                bg='#0a0a0a'
            )
            self.progress_label.pack(padx=30, pady=15)
            
            self.progress_bar = ttk.Progressbar(
                self.progress_window,
                orient=tk.HORIZONTAL,
                length=400,
                mode='indeterminate',
                style='green.Horizontal.TProgressbar'
            )
            self.progress_bar.pack(pady=15)
            
            style = ttk.Style()
            style.theme_use('clam')
            style.configure(
                'green.Horizontal.TProgressbar',
                background='#00ff00',
                troughcolor='#002200'
            )
            
            self.progress_window.grab_set()

    def update_progress(self, value: int, message: str = None):
        """Atualiza o valor da barra de progresso e a mensagem."""
        if self.progress_bar:
            self.progress_bar['value'] = value
            if message and self.progress_label:
                self.progress_label.config(text=message)
            self.progress_window.update_idletasks()

    def hide_progress(self):
        """Fecha a janela de progresso."""
        if self.progress_window:
            self.progress_window.grab_release()
            self.progress_window.destroy()
            self.progress_window = None
            self.progress_bar = None
            self.progress_label = None

    def process_file(self):
        """Inicia o processamento de arquivo com barra de progresso."""
        self.ui_command('disable_buttons')
        self.ui_command('update_status', "Aguardando seleção de arquivo...")
        file_path = filedialog.askopenfilename(
            filetypes=[("Documentos", "*.pdf *.docx *.xlsx *.txt"), ("Todos", "*.*")]
        )
        if file_path:
            self.ui_command('show_progress', "Processando arquivo...")
            threading.Thread(target=self._process_file, args=(file_path,), daemon=True).start()
        else:
            self.ui_command('enable_buttons')
            self.ui_command('update_status', "Operação cancelada")

    def _process_file(self, file_path: str):
        """Simula o processamento de um arquivo com barra de progresso."""
        try:
            total_steps = 100
            for i in range(total_steps + 1):
                time.sleep(0.05)  # Simula um processamento demorado
                progress = int((i / total_steps) * 100)
                self.ui_command('update_progress', progress, f"Processando... {progress}%")
            
            content = self.processor.extract_from_file(file_path)
            self.processor.document_content = content

            self.ui_command('update_status', "Selecione o template...")
            template_path = filedialog.askopenfilename(
                filetypes=[("Templates", "*.docx *.xlsx *.pdf")]
            )
            if template_path:
                self.ui_command('update_status', "Gerando saída...")
                output_path = filedialog.asksaveasfilename(
                    defaultextension=os.path.splitext(template_path)[1]
                )
                if output_path:
                    self.processor.fill_template(template_path, output_path, {"content": content})
                    self.ui_command('show_info', "Arquivo processado com sucesso!")
        except Exception as e:
            self.ui_command('show_error', str(e))
        finally:
            self.ui_command('hide_progress')
            self.ui_command('update_status', "Pronto para novas operações")
            self.ui_command('enable_buttons')

    def setup_ui_handler(self):
        def check_queue():
            try:
                while not self.ui_queue.empty():
                    action, *args = self.ui_queue.get_nowait()
                    handler = {
                        'show_progress': self.show_progress,
                        'update_progress': self.update_progress,
                        'hide_progress': self.hide_progress,
                        'show_error': self.show_error,
                        'show_info': self.show_info,
                        'update_status': lambda x: self.status_var.set(f"» {x}"),
                        'enable_buttons': self.enable_buttons,
                        'disable_buttons': self.disable_buttons
                    }.get(action)
                    
                    if handler:
                        handler(*args) if args else handler()
                        
            except Exception as e:
                print(f"Erro na fila: {str(e)}")
            finally:
                self.root.after(100, check_queue)

        self.root.after(100, check_queue)


    def ui_command(self, action, *args):
        self.ui_queue.put((action, *args))
            
    def hide_progress(self):
        if self.progress_window:
            self.progress_bar.stop()
            self.progress_window.grab_release()
            self.progress_window.destroy()
            self.progress_window = None
            self.progress_bar = None

    def show_error(self, message):
        messagebox.showerror("Erro de Sistema", message)
        self.ui_command('update_status', f"Erro: {message}")

    def show_info(self, message):
        messagebox.showinfo("Operação Concluída", message)
        self.ui_command('update_status', message)

    def enable_buttons(self):
        if self.process_file_btn and self.scrape_btn and self.chat_btn and self.template_btn:
            self.process_file_btn.config(state=tk.NORMAL)
            self.scrape_btn.config(state=tk.NORMAL)
            self.chat_btn.config(state=tk.NORMAL)
            self.template_btn.config(state=tk.NORMAL)

    def disable_buttons(self):
        if self.process_file_btn and self.scrape_btn and self.chat_btn and self.template_btn:
            self.process_file_btn.config(state=tk.DISABLED)
            self.scrape_btn.config(state=tk.DISABLED)
            self.chat_btn.config(state=tk.DISABLED)
            self.template_btn.config(state=tk.DISABLED)

    def process_file(self):
        self.ui_command('disable_buttons')
        self.ui_command('update_status', "Aguardando seleção de arquivo...")
        file_path = filedialog.askopenfilename(
            filetypes=[("Documentos", "*.pdf *.docx *.xlsx *.txt"), ("Todos", "*.*")]
        )
        if file_path:
            self.ui_command('update_status', "Processando arquivo...")
            threading.Thread(target=self._process_file, args=(file_path,), daemon=True).start()
        else:
            self.ui_command('enable_buttons')
            self.ui_command('update_status', "Operação cancelada")

    def _process_file(self, file_path: str):
        try:
            self.ui_command('show_progress', "Decodificando arquivo...")
            content = self.processor.extract_from_file(file_path)
            self.processor.document_content = content

            data = {"content": content}

            self.ui_command('update_status', "Selecione o template...")
            template_path = filedialog.askopenfilename(
                filetypes=[("Templates", "*.docx *.xlsx *.pdf")]
            )
            if template_path:
                self.ui_command('update_status', "Gerando saída...")
                output_path = filedialog.asksaveasfilename(
                    defaultextension=os.path.splitext(template_path)[1]
                )
                if output_path:
                    self.processor.fill_template(template_path, output_path, data)
                    self.ui_command('show_info', "Arquivo processado com sucesso!")
        except Exception as e:
            self.ui_command('show_error', str(e))
        finally:
            self.ui_command('hide_progress')
            self.ui_command('update_status', "Pronto para novas operações")
            self.ui_command('enable_buttons')


    def scrape_website(self):
        try:
            self.ui_command('disable_buttons')
            self.ui_command('update_status', "Aguardando URL...")

            url = simpledialog.askstring("Scraping Web", "Digite a URL do site:")
            if not url:
                self.ui_command('enable_buttons')
                self.ui_command('update_status', "Operação cancelada")
                return

            # Adicionar validação completa
            if not url.startswith(('http://', 'https://')):
                url = f'https://{url}'

            if not validators.url(url):
                raise ValueError("URL inválida")

            self.ui_command('update_status', "Analisando site...")

            # Adicionar headers para evitar bloqueio
            headers = {
                'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36'
            }

            response = requests.get(url, headers=headers, timeout=15)
            response.raise_for_status()

            # Executa o scraping em uma thread separada
            threading.Thread(target=self._scrape_and_process, args=(url,), daemon=True).start()

        except requests.exceptions.RequestException as e:
            self.ui_command('update_status', f"Erro ao acessar o site: {e}")
            self.ui_command('enable_buttons')
        except ValueError as e:
            self.ui_command('update_status', str(e))
            self.ui_command('enable_buttons')


    def _scrape_and_process(self, url: str):
        try:
            self.ui_command('show_progress', "Coletando dados do site...")
            data = self.processor.scrape_website(url)
            self.processor.scraped_data = data

            self.ui_command('update_status', "Selecione o template...")
            template_path = filedialog.askopenfilename(
                filetypes=[("Templates", "*.docx *.xlsx *.pdf")]
            )
            if template_path:
                self.ui_command('update_status', "Gerando relatório...")
                output_path = filedialog.asksaveasfilename(
                    defaultextension=os.path.splitext(template_path)[1]
                )
                if output_path:
                    self.processor.fill_template(template_path, output_path, data)
                    self.ui_command('show_info', "Dados processados com sucesso!")
        except Exception as e:
            self.ui_queue.put(('show_error', str(e)))
        finally:
            self.ui_queue.put(('hide_progress',))
            self.ui_queue.put(('update_status', "Pronto para novas operações"))
            self.ui_queue.put(('enable_buttons',))

    def create_template(self):
        self.ui_command('disable_buttons')
        template_type = simpledialog.askstring(
            "Criar Template",
            "Digite o tipo de template (docx/xlsx) e campos separados por vírgula:\n"
            "Ex: docx, nome, data, valor"
        )
        
        if template_type:
            parts = [p.strip() for p in template_type.split(',')]
            threading.Thread(
                target=self._generate_template_thread,
                args=(parts[0], parts[1:]),
                daemon=True
            ).start()

    def _generate_template_thread(self, template_type: str, fields: list):
        try:
            self.ui_command('show_progress', "Gerando template com IA...")
            output_path = self.processor.generate_template(template_type, fields)
            self.ui_command('show_info', f"Template criado: {output_path}")
        except Exception as e:
            self.ui_command('show_error', str(e))
        finally:
            self.ui_command('hide_progress')
            self.ui_command('enable_buttons')

    def chat_with_ai(self):
        chat_window = tk.Toplevel(self.root)
        chat_window.title("Assistente IA")
        chat_window.geometry("800x600")
        chat_window.configure(bg=Config.CHAT_BG)

        main_frame = tk.Frame(chat_window, bg=Config.CHAT_BG)
        main_frame.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)

        self.chat_text = tk.Text(
            main_frame,
            wrap=tk.WORD,
            bg=Config.CHAT_BG,
            fg=Config.TEXT_COLOR,
            font=('Segoe UI', 12),
            state=tk.DISABLED
        )
        self.chat_text.pack(fill=tk.BOTH, expand=True)

        input_frame = tk.Frame(main_frame, bg=Config.CHAT_BG)
        input_frame.pack(fill=tk.X, pady=10)

        self.input_entry = tk.Entry(
            input_frame,
            bg=Config.SECONDARY_COLOR,
            fg=Config.TEXT_COLOR,
            font=('Segoe UI', 12),
            width=50
        )
        self.input_entry.pack(side=tk.LEFT, fill=tk.X, expand=True, padx=5)
        self.input_entry.bind("<Return>", lambda e: self.process_query(chat_window))

        tk.Button(
            input_frame,
            text="Enviar",
            command=lambda: self.process_query(chat_window),
            bg=Config.ACCENT_COLOR,
            fg=Config.TEXT_COLOR,
            width=10
        ).pack(side=tk.RIGHT)

        self.update_chat("IA: Bem-vindo ao DataScan!\nComo posso ajudá-lo hoje?\n", "system")

        
    def upload_file_for_chat(self, parent_window):
        file_path = filedialog.askopenfilename(
            filetypes=[("Documentos", "*.pdf *.docx *.xlsx *.txt"), ("Todos", "*.*")]
        )
        if file_path:
            try:
                content = self.processor.extract_from_file(file_path)
                self.processor.document_content = content
                self.update_chat(f"Sistema: Arquivo carregado com sucesso: {os.path.basename(file_path)}\n", "system")
                self.update_chat(f"Resumo do documento:\n{content[:500]}...\n\n", "file")
            except Exception as e:
                self.update_chat(f"Erro ao processar arquivo: {str(e)}\n", "error")

    def clear_chat_context(self):
        self.processor.chat_context = []
        self.update_chat("Sistema: Contexto da conversa limpo com sucesso\n", "system")

    def process_query(self, window):
        query = self.input_entry.get().strip()
        self.input_entry.delete(0, tk.END)
        if not query:
            return
            
        self.update_chat(f"Você: {query}\n", "user")
        threading.Thread(target=self.generate_response, args=(query, window), daemon=True).start()


    def handle_template_creation(self, query: str):
        try:
            self.update_chat("Assistente: Analisando sua solicitação...\n", "system")
            
            prompt = (
                f"Com base na seguinte solicitação: '{query}'\n"
                "Identifique:\n"
                "1. Tipo de documento (docx/xlsx)\n"
                "2. Campos necessários\n"
                "3. Formatação especial\n"
                "Retorne no formato JSON:"
                '{"type": "docx", "fields": ["nome", "data"], "instructions": "..."}'
            )
            
            response = self.processor.generate_ai_response(prompt)
            params = json.loads(response)
            
            self.update_chat(
                f"Assistente: Criando template {params['type'].upper()} com campos: "
                f"{', '.join(params['fields'])}\n",
                "system"
            )
            
            template_path = self.processor.generate_template(
                params['type'],
                params['fields']
            )
            
            self.update_chat(
                f"✅ Template criado com sucesso!\n"
                f"📂 Caminho: {template_path}\n"
                f"📝 Instruções: {params.get('instructions', 'Nenhuma')}\n\n",
                "success"
            )
            
        except Exception as e:
            self.update_chat(f"❌ Erro na criação: {str(e)}\n", "error")


    def generate_response(self, query, window):
        try:
            start_time = time.time()
            response = self.processor.generate_ai_response(query)
            elapsed_time = time.time() - start_time
            
            # Formatar resposta corretamente
            formatted_response = f"IA ({elapsed_time:.2f}s):\n{response}\n"
            formatted_response += "―" * 50 + "\n"
            
            self.update_chat(formatted_response, "ai")
        except Exception as e:
            self.update_chat(f"Erro: {str(e)}\n", "error")
            
            
    def show_help(self):
        help_text = """Comandos disponíveis:
/ajuda - Mostra esta mensagem
/limpar - Limpa o contexto da conversa
/sair - Fecha o chat
/carregar - Abre diálogo para carregar arquivo
/criar template - Cria um novo template com IA
"""
        self.update_chat(f"Sistema:\n{help_text}\n", "system")

    def update_chat(self, message: str, msg_type: str) -> None:
        color_map = {
            "user": Config.TEXT_COLOR,
            "system": "#B0BEC5",
            "error": "#EF5350",
            "ai": Config.ACCENT_COLOR,
            "success": "#66BB6A",
            "file": "#42A5F5"
        }
        
        self.chat_text.config(state=tk.NORMAL)
        self.chat_text.tag_configure(msg_type, foreground=color_map.get(msg_type, "#FFFFFF"))
        self.chat_text.insert(tk.END, message, msg_type)
        self.chat_text.see(tk.END)
        self.chat_text.config(state=tk.DISABLED)


    def setup_db_explorer(self):
        """Janela para explorar a estrutura do banco de dados"""
        self.explorer_window = tk.Toplevel(self.root)
        self.explorer_window.title("Explorador de Banco de Dados")
        self.explorer_window.geometry("600x400")
        
        self.tree = ttk.Treeview(self.explorer_window)
        self.tree.pack(fill=tk.BOTH, expand=True)
        
        vsb = ttk.Scrollbar(self.explorer_window, orient="vertical", command=self.tree.yview)
        vsb.pack(side=tk.RIGHT, fill=tk.Y)
        self.tree.configure(yscrollcommand=vsb.set)
        
        self.explorer_window.withdraw()

    def show_db_explorer(self):
        """Atualiza e exibe o explorador de banco de dados"""
        try:
            schema = self.processor.db_manager.get_schema()
            self.tree.delete(*self.tree.get_children())
            
            for table in schema['tables']:
                table_id = self.tree.insert("", "end", text=table['name'], values=["Table"])
                for column in table['columns']:
                    self.tree.insert(table_id, "end", 
                                   text=f"{column['name']} ({column['type']})",
                                   values=["Column"])
            
            self.explorer_window.deiconify()
        except Exception as e:
            self.show_error(str(e))

    def create_backup(self):
        """Inicia processo de backup do banco de dados"""
        try:
            backup_path = self.processor.db_manager.backup_database()
            self.show_info(f"Backup criado com sucesso em:\n{backup_path}")
        except Exception as e:
            self.show_error(str(e))

    def show_chart(self, data: dict):
        """Exibe os dados em formato gráfico"""
        if self.current_chart_window:
            self.current_chart_window.destroy()
            
        self.current_chart_window = tk.Toplevel(self.root)
        self.current_chart_window.title("Visualização de Dados")
        
        fig = plt.Figure(figsize=(6, 4), dpi=100)
        ax = fig.add_subplot(111)
        
        if isinstance(data, list) and len(data) > 0:
            df = pd.DataFrame(data)
            df.plot(kind='bar', ax=ax)
            
            canvas = FigureCanvasTkAgg(fig, master=self.current_chart_window)
            canvas.draw()
            canvas.get_tk_widget().pack(fill=tk.BOTH, expand=True)



if __name__ == "__main__":
    if not os.path.exists(Config.TEMP_DIR):
        os.makedirs(Config.TEMP_DIR)
    if not os.path.exists(Config.BACKUP_DIR):
        os.makedirs(Config.BACKUP_DIR)
        
    app = Application()
    app.root.mainloop()