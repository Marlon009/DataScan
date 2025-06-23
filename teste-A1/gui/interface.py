from PyQt6.QtWidgets import (
    QMainWindow, QWidget, QVBoxLayout, QHBoxLayout,
    QTabWidget, QLineEdit, QPushButton, QTextEdit,
    QComboBox, QLabel, QFileDialog, QMessageBox, QMenuBar, QMenu
)
from PyQt6.QtCore import Qt, QEvent
from PyQt6.QtGui import QAction, QIcon, QPalette, QColor
from services.file_processor import process_file
from services.web_scraper import scrape_website
from services.ai_terminal import AITerminal
import os
import sys

class DataScan(QMainWindow):
    def __init__(self):
        super().__init__()
        self.setWindowTitle("DataScan Pro")
        self.setGeometry(100, 100, 1200, 800)
        self.ai = AITerminal()
        self.dark_mode = False
        self.current_file_path = None
        
        # Configurar ícone da janela
        try:
            self.setWindowIcon(QIcon("icon.png"))
        except:
            pass
            
        self.init_ui()
        
    def keyPressEvent(self, event):
        if event.key() == Qt.Key.Key_Return and not event.modifiers():
            current_tab = self.tabs.currentIndex()
            
            if current_tab == 1:  # Aba Web
                self.scrape_web()
            elif current_tab == 2:  # Aba IA
                self.handle_ai_input()
        else:
            super().keyPressEvent(event)

    def init_ui(self):
        # Criar barra de menu
        self.create_menu_bar()
        
        main_widget = QWidget()
        self.setCentralWidget(main_widget)
        layout = QVBoxLayout(main_widget)
        layout.setContentsMargins(10, 10, 10, 10)

        self.tabs = QTabWidget()
        self.tabs.setTabPosition(QTabWidget.TabPosition.North)
        self.tabs.setMovable(True)
        layout.addWidget(self.tabs)

        # File Processing Tab
        file_tab = QWidget()
        self.init_file_tab(file_tab)
        self.tabs.addTab(file_tab, self.create_tab_icon("📁"), "Arquivos")

        # Web Scraping Tab
        web_tab = QWidget()
        self.init_web_tab(web_tab)
        self.tabs.addTab(web_tab, self.create_tab_icon("🌐"), "Web")

        # AI Terminal Tab
        ai_tab = QWidget()
        self.init_ai_tab(ai_tab)
        self.tabs.addTab(ai_tab, self.create_tab_icon("🤖"), "IA")

    def init_file_tab(self, tab):
        layout = QVBoxLayout(tab)
        
        # File Selection
        file_layout = QHBoxLayout()
        self.file_path = QLineEdit()
        btn_browse = QPushButton("Selecionar Arquivo")
        btn_browse.clicked.connect(self.browse_file)
        
        file_layout.addWidget(self.file_path)
        file_layout.addWidget(btn_browse)
        
        # Preview
        self.preview = QTextEdit()
        
        # Processing
        btn_process = QPushButton("Processar")
        btn_process.clicked.connect(self.process_file)
        
        layout.addLayout(file_layout)
        layout.addWidget(self.preview)
        layout.addWidget(btn_process)
        
    def create_tab_icon(self, emoji):
        label = QLabel(emoji)
        label.setAlignment(Qt.AlignmentFlag.AlignCenter)
        return label

    def init_web_tab(self, tab):
        layout = QVBoxLayout(tab)
        
        # URL Input
        url_layout = QHBoxLayout()
        self.url_input = QLineEdit()
        btn_scrape = QPushButton("Escanear")
        btn_scrape.clicked.connect(self.scrape_web)
        
        url_layout.addWidget(QLabel("URL:"))
        url_layout.addWidget(self.url_input)
        url_layout.addWidget(btn_scrape)
        
        # Results
        self.web_results = QTextEdit()
        
        # Export Buttons
        export_layout = QHBoxLayout()
        btn_export_pdf = QPushButton("Exportar PDF")
        btn_export_excel = QPushButton("Exportar Excel")
        btn_export_docx = QPushButton("Exportar DOCX")
        
        btn_export_pdf.clicked.connect(lambda: self.export_web_results('pdf'))
        btn_export_excel.clicked.connect(lambda: self.export_web_results('excel'))
        btn_export_docx.clicked.connect(lambda: self.export_web_results('docx'))
        
        export_layout.addWidget(btn_export_pdf)
        export_layout.addWidget(btn_export_excel)
        export_layout.addWidget(btn_export_docx)
        
        layout.addLayout(url_layout)
        layout.addWidget(self.web_results)
        layout.addLayout(export_layout)

    def init_ai_tab(self, tab):
        layout = QVBoxLayout(tab)
        
        # API Key Section
        key_layout = QHBoxLayout()
        self.api_key_input = QLineEdit()
        self.api_key_input.setPlaceholderText("Cole sua API Key do Gemini aqui...")
        btn_save_key = QPushButton("🔑 Salvar")
        btn_save_key.clicked.connect(self.save_api_key)
        
        key_layout.addWidget(self.api_key_input)
        key_layout.addWidget(btn_save_key)
        
        # Chat history
        self.chat_history = QTextEdit()
        self.chat_history.setReadOnly(True)
        
        # User input
        input_layout = QHBoxLayout()
        self.ai_input_ia = QLineEdit()
        self.ai_input_ia.setPlaceholderText("Digite sua solicitação...")
        btn_send = QPushButton("🚀 Enviar")
        btn_send.clicked.connect(self.handle_ai_input)
        
        input_layout.addWidget(self.ai_input_ia)
        input_layout.addWidget(btn_send)
        
        layout.addLayout(key_layout)
        layout.addWidget(self.chat_history)
        layout.addLayout(input_layout)

    def save_api_key(self):
        key = self.api_key_input.text().strip()
        if key:
            self.ai = AITerminal(api_key=key)
            QMessageBox.information(self, "Sucesso", "Chave configurada com sucesso!")
        else:
            QMessageBox.warning(self, "Aviso", "Insira uma API Key válida!")

    def create_menu_bar(self):
        menubar = self.menuBar()
        
        # Menu Arquivo
        file_menu = menubar.addMenu("Arquivo")
        
        new_action = QAction("Novo", self)
        open_action = QAction("Abrir...", self)
        save_action = QAction("Salvar", self)
        exit_action = QAction("Sair", self)
        
        file_menu.addAction(new_action)
        file_menu.addAction(open_action)
        file_menu.addAction(save_action)
        file_menu.addSeparator()
        file_menu.addAction(exit_action)
        
        # Menu Visualização
        view_menu = menubar.addMenu("Visualização")
        
        self.dark_mode_action = QAction("Modo Noturno", self, checkable=True)
        self.dark_mode_action.triggered.connect(self.toggle_dark_mode)
        view_menu.addAction(self.dark_mode_action)
        
        # Menu Ajuda
        help_menu = menubar.addMenu("Ajuda")
        about_action = QAction("Sobre", self)
        help_menu.addAction(about_action)

    def toggle_dark_mode(self):
        self.dark_mode = not self.dark_mode
        if self.dark_mode:
            self.apply_dark_theme()
        else:
            self.apply_light_theme()

    def apply_dark_theme(self):
        dark_palette = QPalette()
        
        # Configuração da paleta de cores
        dark_palette.setColor(QPalette.ColorRole.Window, QColor(53, 53, 53))
        dark_palette.setColor(QPalette.ColorRole.WindowText, Qt.GlobalColor.white)
        dark_palette.setColor(QPalette.ColorRole.Base, QColor(35, 35, 35))
        dark_palette.setColor(QPalette.ColorRole.AlternateBase, QColor(53, 53, 53))
        dark_palette.setColor(QPalette.ColorRole.ToolTipBase, Qt.GlobalColor.white)
        dark_palette.setColor(QPalette.ColorRole.ToolTipText, Qt.GlobalColor.white)
        dark_palette.setColor(QPalette.ColorRole.Text, Qt.GlobalColor.white)
        dark_palette.setColor(QPalette.ColorRole.Button, QColor(53, 53, 53))
        dark_palette.setColor(QPalette.ColorRole.ButtonText, Qt.GlobalColor.white)
        dark_palette.setColor(QPalette.ColorRole.BrightText, Qt.GlobalColor.red)
        dark_palette.setColor(QPalette.ColorRole.Link, QColor(42, 130, 218))
        dark_palette.setColor(QPalette.ColorRole.Highlight, QColor(42, 130, 218))
        dark_palette.setColor(QPalette.ColorRole.HighlightedText, Qt.GlobalColor.black)
        
        self.setPalette(dark_palette)
        
        self.setStyleSheet("""
            QTextEdit, QLineEdit, QComboBox {
                background-color: #252525;
                color: #ffffff;
                border: 2px solid #444;
                border-radius: 8px;
                padding: 12px;
                font-size: 14px;
                selection-background-color: #3a3a3a;
            }
            QPushButton {
                background-color: #0078d4;
                color: white;
                border-radius: 6px;
                padding: 10px 20px;
                font-weight: bold;
                min-width: 80px;
            }
            QPushButton:hover {
                background-color: #006cbd;
            }
            QTabWidget::pane {
                border: 1px solid #444;
                background: #353535;
            }
            QTabBar::tab {
                background: #353535;
                color: white;
                padding: 8px;
                border-top-left-radius: 4px;
                border-top-right-radius: 4px;
            }
            QTabBar::tab:selected {
                background: #252525;
                border-bottom-color: #0078d4;
            }
            QMenuBar {
                background-color: #353535;
                color: white;
            }
            QMenuBar::item {
                background-color: transparent;
                padding: 5px 10px;
            }
            QMenuBar::item:selected {
                background-color: #555;
            }
            QMenu {
                background-color: #353535;
                color: white;
                border: 1px solid #444;
            }
            QMenu::item:selected {
                background-color: #0078d4;
            }
        """)

    def apply_light_theme(self):
        self.setPalette(self.style().standardPalette())
        
        self.setStyleSheet("""
            QTextEdit, QLineEdit, QComboBox {
                background-color: #ffffff;
                color: #333333;
                border: 2px solid #0078d4;
                border-radius: 8px;
                padding: 12px;
                font-size: 14px;
                selection-background-color: #b8d8ff;
            }
            QPushButton {
                background-color: #0078d4;
                color: white;
                border-radius: 6px;
                padding: 10px 20px;
                font-weight: bold;
                min-width: 80px;
            }
            QPushButton:hover {
                background-color: #006cbd;
            }
            QTabWidget::pane {
                border: 1px solid #ddd;
                background: #f0f0f0;
            }
            QTabBar::tab {
                background: #f0f0f0;
                color: #333;
                padding: 8px;
                border-top-left-radius: 4px;
                border-top-right-radius: 4px;
            }
            QTabBar::tab:selected {
                background: #ffffff;
                border-bottom-color: #0078d4;
            }
            QMenuBar {
                background-color: #f0f0f0;
                color: #333;
            }
            QMenuBar::item:selected {
                background-color: #e0e0e0;
            }
            QMenu {
                background-color: #ffffff;
                color: #333;
                border: 1px solid #ddd;
            }
            QMenu::item:selected {
                background-color: #0078d4;
                color: white;
            }
        """)

    def browse_file(self):
        path, _ = QFileDialog.getOpenFileName(
            self, "Selecionar Arquivo", "", 
            "Documentos (*.docx *.pdf *.xlsx *.xls *.csv)"
        )
        if path:
            self.file_path.setText(path)
            self.preview.setText(f"Arquivo selecionado: {path}")

    def process_file(self):
        path = self.file_path.text()
        if path:
            result = process_file(path)
            self.preview.setText(result)
            QMessageBox.information(self, "Sucesso", "Arquivo processado!")

    def scrape_web(self):
        url = self.url_input.text()
        if url:
            result = scrape_website(url)
            self.web_results.setText(result)

    def handle_ai_input(self):
        user_text = self.ai_input_ia.text()
        if not user_text:
            return
            
        self.chat_history.append(f"Você: {user_text}")
        self.ai_input_ia.clear()
        
        response, file_path = self.ai.generate_response(user_text)
        
        if file_path:
            self.current_file_path = file_path
            self.chat_history.append(f"IA: {response}\n[Abrir Documento]({file_path})")
        else:
            self.chat_history.append(f"IA: {response}")
        
        # Auto-scroll to latest message
        self.chat_history.verticalScrollBar().setValue(
            self.chat_history.verticalScrollBar().maximum()
        )
    
    def export_web_results(self, format):
        content = self.web_results.toPlainText()
        if not content:
            QMessageBox.warning(self, "Aviso", "Nenhum resultado para exportar!")
            return
        
        filename, _ = QFileDialog.getSaveFileName(
            self, "Salvar Resultado", "", 
            f"{format.upper()} Files (*.{format})"
        )
        
        if not filename:
            return
        
        try:
            if format == 'pdf':
                from fpdf import FPDF
                pdf = FPDF()
                pdf.add_page()
                pdf.set_font("Arial", size=12)
                pdf.multi_cell(0, 10, content)
                pdf.output(filename)
                
            elif format == 'excel':
                import pandas as pd
                data = []
                current_item = {}
                for line in content.split('\n'):
                    if line.startswith('País:'):
                        if current_item:
                            data.append(current_item)
                        current_item = {'País': line.split(':')[1].strip()}
                    elif ':' in line:
                        key, value = line.split(':', 1)
                        current_item[key.strip()] = value.strip()
                
                if current_item:
                    data.append(current_item)
                    
                df = pd.DataFrame(data)
                df.to_excel(filename, index=False)
                
            elif format == 'docx':
                from docx import Document
                doc = Document()
                doc.add_heading('Resultado da Web Scraping', 0)
                doc.add_paragraph(content)
                doc.save(filename)
                
            QMessageBox.information(self, "Sucesso", f"Arquivo salvo em:\n{filename}")
            
        except Exception as e:
            QMessageBox.critical(self, "Erro", f"Falha ao exportar:\n{str(e)}")

if __name__ == "__main__":
    from PyQt6.QtWidgets import QApplication
    app = QApplication([])
    
    # Configuração para alta DPI (melhor visualização em monitores 4K)
    app.setAttribute(Qt.ApplicationAttribute.AA_EnableHighDpiScaling)
    app.setAttribute(Qt.ApplicationAttribute.AA_UseHighDpiPixmaps)
    
    window = DataScan()
    window.show()
    app.exec()