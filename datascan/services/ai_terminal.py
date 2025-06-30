import google.generativeai as genai
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from fpdf import FPDF
import pandas as pd
import os
import re
import time
import sys
from dotenv import load_dotenv
from collections import defaultdict

load_dotenv()

class AITerminal:
    def __init__(self, api_key=None):
        self.api_key = api_key or os.getenv('GEMINI_API_KEY')
        self.model = None
        self.conversation_state = {
            'awaiting_response': False,
            'document_type': None,
            'fields': [],
            'content': None,
            'formatting': {}
        }
        self.initialize_model()
        
    def initialize_model(self):
        if self.api_key:
            try:
                genai.configure(api_key=self.api_key)
                self.model = genai.GenerativeModel('gemini-1.5-flash')
                return True
            except Exception as e:
                print(f"Erro na inicialização: {e}")
                return False
        return False

    def generate_response(self, prompt):
        if not self.model:
            return "Erro: Configure sua API Key primeiro!", None
            
        try:
            if not self.conversation_state['awaiting_response']:
                self._reset_conversation()
                return self._init_document_creation(prompt)
            else:
                return self._handle_document_details(prompt)
                
        except Exception as e:
            return f"Erro na geração: {str(e)}", None

    def _init_document_creation(self, prompt):
        doc_type = self._detect_document_type(prompt)
        if doc_type:
            self.conversation_state.update({
                'awaiting_response': True,
                'document_type': doc_type
            })
            return (
                f"Vou ajudar com o {doc_type.upper()}. Por favor, informe:\n"
                "1. Quais campos devem ser incluídos (separados por vírgula)\n"
                "2. Algum conteúdo base específico?\n"
                "3. Formatação especial (tabelas, cores, etc.)\n"
                "Exemplo: 'Nome, Data, Hora | Texto base: Contrato | Tabela com 3 colunas'",
                None
            )
        else:
            response = self.model.generate_content(prompt)
            return response.text, None

    def _handle_document_details(self, user_input):
        try:
            details = [part.strip() for part in user_input.split('|')]
            
            if len(details) > 0:
                self.conversation_state['fields'] = [f.strip() for f in details[0].split(',')]
                
            if len(details) > 1:
                content_part = details[1]
                if ':' in content_part:
                    content_part = content_part.split(':', 1)[1].strip()
                self.conversation_state['content'] = content_part
                
            if len(details) > 2:
                self._process_formatting(details[2])
                
            if not self.conversation_state['fields']:
                return "Erro: Por favor, informe pelo menos um campo.", None
                
            if not self.conversation_state['document_type']:
                return "Erro: Tipo de documento não especificado.", None
                
            file_path = self._create_document()
            response = f"Documento gerado com sucesso!\nLocal: {file_path}"
            return response, file_path
            
        except Exception as e:
            return f"Erro na geração do documento: {str(e)}", None
        finally:
            self._reset_conversation()  # Reset SEMPRE após processamento

    def _create_document(self):
        doc_type = self.conversation_state['document_type']
        creators = {
            'docx': self._create_custom_docx,
            'pdf': self._create_custom_pdf,
            'excel': self._create_custom_excel
        }
        
        creator = creators.get(doc_type)
        if creator:
            return creator()
        else:
            return "Tipo de documento não suportado"

    def _create_custom_docx(self):
        try:
            doc = Document()
            
            # Adiciona título com estilo aprimorado
            if self.conversation_state['content']:
                title = doc.add_paragraph()
                title_run = title.add_run(self.conversation_state['content'])
                title_run.bold = True
                title_run.font.size = Pt(16)
                title_run.font.color.rgb = RGBColor(0, 0, 100)  # Azul escuro
                title.alignment = WD_ALIGN_PARAGRAPH.CENTER
                doc.add_paragraph()  # Espaçamento
            
            # Adiciona campos com formatação
            for field in self.conversation_state['fields']:
                p = doc.add_paragraph()
                field_run = p.add_run(f"{field}:")
                field_run.bold = True
                p.add_run(" " * 5)  # Espaço
                p.add_run("_________________________")
            
            # Adiciona tabela se solicitado
            if 'table' in self.conversation_state['formatting']:
                cols = self.conversation_state['formatting']['table'].get('columns', 2)
                rows = self.conversation_state['formatting']['table'].get('rows', 1)
                table = doc.add_table(rows=rows, cols=cols)
                table.style = 'Table Grid'
                
                # Preenche a primeira linha com cabeçalhos
                if self.conversation_state['fields']:
                    header_cells = table.rows[0].cells
                    for i in range(min(cols, len(self.conversation_state['fields']))):
                        header_cells[i].text = self.conversation_state['fields'][i]
                        header_cells[i].paragraphs[0].runs[0].bold = True
            
            # Gera nome de arquivo único e seguro
            filename = f"documento_{int(time.time())}.docx"
            doc.save(filename)
            return os.path.abspath(filename)
            
        except Exception as e:
            return f"Erro ao criar DOCX: {str(e)}"

    def _create_custom_pdf(self):
        try:
            # Gera nome de arquivo único
            filename = f"documento_{int(time.time())}.pdf"
            pdf = FPDF()
            pdf.add_page()
            
            # Título com formatação
            if self.conversation_state['content']:
                pdf.set_font("Arial", 'B', 16)
                pdf.set_text_color(0, 0, 100)  # Azul escuro
                pdf.cell(0, 15, txt=self.conversation_state['content'], ln=1, align='C')
                pdf.ln(5)
            
            # Campos
            pdf.set_font("Arial", size=12)
            pdf.set_text_color(0, 0, 0)  # Preto
            for field in self.conversation_state['fields']:
                pdf.cell(50, 10, txt=f"{field}:", ln=0)
                pdf.set_font("Arial", 'U', 12)  # Sublinhado para área de preenchimento
                pdf.cell(0, 10, txt="_________________________", ln=1)
                pdf.set_font("Arial", size=12)
                pdf.ln(5)
            
            # Tabela
            if 'table' in self.conversation_state['formatting']:
                cols = self.conversation_state['formatting']['table'].get('columns', 2)
                rows = self.conversation_state['formatting']['table'].get('rows', 1)
                col_width = 190 / cols
                
                # Cabeçalhos
                pdf.set_font("Arial", 'B', 12)
                for i in range(min(cols, len(self.conversation_state['fields']))):
                    pdf.cell(col_width, 10, self.conversation_state['fields'][i], border=1, align='C')
                pdf.ln()
                
                # Linhas
                pdf.set_font("Arial", size=12)
                for _ in range(rows):
                    for col in range(cols):
                        pdf.cell(col_width, 10, "", border=1)
                    pdf.ln()
            
            pdf.output(filename)
            return os.path.abspath(filename)
            
        except Exception as e:
            return f"Erro ao criar PDF: {str(e)}"

    def _create_custom_excel(self):
        try:
            # Gera nome de arquivo único
            filename = f"documento_{int(time.time())}.xlsx"
            
            # Cria DataFrame com validação
            if self.conversation_state['fields']:
                df = pd.DataFrame(columns=self.conversation_state['fields'])
                
                # Adiciona conteúdo base
                if self.conversation_state['content']:
                    base_content = {
                        self.conversation_state['fields'][0]: self.conversation_state['content']
                    }
                    df = pd.DataFrame([base_content])
                
                # Adiciona linhas extras se solicitado
                if 'table' in self.conversation_state['formatting']:
                    rows = self.conversation_state['formatting']['table'].get('rows', 1)
                    for _ in range(rows - 1):
                        df = pd.concat([df, pd.DataFrame([{}])], ignore_index=True)
                
                df.to_excel(filename, index=False)
                return os.path.abspath(filename)
            else:
                return "Erro: Nenhum campo definido para a planilha"
                
        except Exception as e:
            return f"Erro ao criar Excel: {str(e)}"

    def _reset_conversation(self):
        self.conversation_state = {
            'awaiting_response': False,
            'document_type': None,
            'fields': [],
            'content': None,
            'formatting': {}
        }
        
    def _detect_document_type(self, text):
        text = text.lower()
        doc_types = {
            'docx': ['docx', 'word', 'documento'],
            'pdf': ['pdf', 'relatório', 'formulário'],
            'excel': ['excel', 'planilha', 'xlsx', 'csv', 'tabela']
        }
        
        for doc_type, keywords in doc_types.items():
            if any(keyword in text for keyword in keywords):
                return doc_type
        return None

    def _process_formatting(self, formatting_text):
        formatting_text = formatting_text.lower()
        
        # Processa tabela
        if 'tabela' in formatting_text:
            cols = 2
            rows = 1
            col_match = re.search(r'(\d+)\s*colunas?', formatting_text)
            if col_match:
                cols = max(1, min(10, int(col_match.group(1))))  # Limita entre 1-10 colunas
            row_match = re.search(r'(\d+)\s*linhas?', formatting_text)
            if row_match:
                rows = max(1, min(50, int(row_match.group(1))))  # Limita entre 1-50 linhas
            self.conversation_state['formatting']['table'] = {'columns': cols, 'rows': rows}
        
        # Processa cores
        color_match = re.search(r'cor\s*([a-záéíóúãõâêôç]+)', formatting_text)
        if color_match:
            self.conversation_state['formatting']['color'] = color_match.group(1)
        
        # Processa fonte
        font_match = re.search(r'fonte\s*([a-záéíóúãõâêôç\s]+)', formatting_text)
        if font_match:
            self.conversation_state['formatting']['font'] = font_match.group(1).strip()

    def interactive_mode(self):
        print("Bem-vindo ao assistente de documentos! (Digite 'sair' para encerrar)")
        while True:
            user_input = input("\n> ").strip()
            if user_input.lower() in ['sair', 'exit', 'quit']:
                break
                
            response, filepath = self.generate_response(user_input)
            print(f"\n{response}")
            
            if filepath and os.path.exists(filepath):
                print(f"\nArquivo criado em: {filepath}")
                open_file = input("Deseja abrir o arquivo? [s/n] ").strip().lower()
                if open_file == 's':
                    try:
                        if sys.platform == "win32":
                            os.startfile(filepath)
                        elif sys.platform == "darwin":
                            os.system(f'open "{filepath}"')
                        else:
                            os.system(f'xdg-open "{filepath}"')
                    except Exception as e:
                        print(f"Erro ao abrir arquivo: {str(e)}")
            elif filepath and not os.path.exists(filepath):
                print(f"Erro: O arquivo não foi criado - {filepath}")

# Exemplo de uso
if __name__ == "__main__":
    terminal = AITerminal()
    if terminal.model:
        terminal.interactive_mode()
    else:
        print("Não foi possível inicializar o modelo Gemini. Verifique sua API KEY.")