from PyQt6.QtWidgets import (
    QMainWindow, QWidget, QVBoxLayout, QHBoxLayout,
    QTabWidget, QLineEdit, QPushButton, QTextEdit,
    QComboBox, QLabel, QFileDialog, QMessageBox, QMenuBar, QMenu, QTextBrowser,
    QDialogButtonBox, QDialog, QSizePolicy, QTextBrowser
)
from PyQt6.QtCore import Qt, QEvent
from PyQt6.QtGui import QAction, QIcon, QPalette, QColor
from services.file_processor import process_file
from services.web_scraper import scrape_website
from services.ai_terminal import AITerminal
import os
import sys

class DataScan(QMainWindow):
    def __init__(self):
        super().__init__()
        self.setWindowTitle("DataScan")
        self.setWindowIcon(QIcon("datascan.ico"))

        self.setGeometry(100, 100, 1200, 800)
        self.ai = AITerminal()
        self.dark_mode = False
        self.current_file_path = None
            
        self.init_ui()
        self.apply_light_theme()  # Tema claro por padrão
        
    def init_ui(self):
        # Criar barra de menu
        self.create_menu_bar()
        
        main_widget = QWidget()
        self.setCentralWidget(main_widget)
        layout = QVBoxLayout(main_widget)
        layout.setContentsMargins(10, 10, 10, 10)

        self.tabs = QTabWidget()
        self.tabs.setTabPosition(QTabWidget.TabPosition.North)
        self.tabs.setMovable(True)
        layout.addWidget(self.tabs)

        # File Processing Tab
        file_tab = QWidget()
        self.init_file_tab(file_tab)
        self.tabs.addTab(file_tab, "📁 Arquivos")

        # Web Scraping Tab
        web_tab = QWidget()
        self.init_web_tab(web_tab)
        self.tabs.addTab(web_tab, "🌐 Web")

        # AI Terminal Tab
        ai_tab = QWidget()
        self.init_ai_tab(ai_tab)
        self.tabs.addTab(ai_tab, "🤖 IA")

    def create_menu_bar(self):
        menubar = self.menuBar()
        
        # Menu Arquivo
        
        # Menu Visualização
        view_menu = menubar.addMenu("Visualização")
        
        self.dark_mode_action = QAction("Modo Noturno", self, checkable=True)
        self.dark_mode_action.triggered.connect(self.toggle_dark_mode)
        view_menu.addAction(self.dark_mode_action)
        
        # Menu Ajuda
        help_menu = menubar.addMenu("Ajuda")
        about_action = QAction("Sobre", self)
        about_action.triggered.connect(self.show_tutorial)
        help_menu.addAction(about_action)
        help_menu.addAction(about_action)
        
    def show_tutorial(self):
        dialog = QDialog(self)
        dialog.setWindowTitle("📘 Tutorial e Informações")
        dialog.setMinimumSize(820, 600)

        tab_widget = QTabWidget(dialog)

        # === ABA: Sobre o Programa ===
        about_tab = QWidget()
        about_layout = QVBoxLayout(about_tab)
        about_text = QTextBrowser()
        about_text.setOpenExternalLinks(True)
        about_text.setSizePolicy(QSizePolicy.Policy.Expanding, QSizePolicy.Policy.Expanding)

        about_html = """
        <h1>DataScan</h1>
        <p><b>Versão:</b> 1.0</p>
        <p>O <b>DataScan</b> é uma ferramenta multifuncional que combina:</p>
        <ul>
            <li>📄 Processamento inteligente de documentos (DOCX, PDF, Excel)</li>
            <li>🌐 Web scraping avançado para extração de dados</li>
            <li>🤖 Assistente de IA para geração de conteúdos</li>
        </ul>
        <p>Desenvolvido por <b>Marlon</b> —
        <a href='https://github.com/Marlon009'>GitHub</a></p>
        """
        about_text.setHtml(about_html)
        about_layout.addWidget(about_text)
        tab_widget.addTab(about_tab, "Sobre")

        # === ABA: Tutorial ===
        tutorial_tab = QWidget()
        tutorial_layout = QVBoxLayout(tutorial_tab)
        tutorial_text = QTextBrowser()
        tutorial_text.setOpenExternalLinks(True)
        tutorial_text.setSizePolicy(QSizePolicy.Policy.Expanding, QSizePolicy.Policy.Expanding)

        tutorial_html = """
        <h2>📁 Aba Arquivos</h2>
        <ol>
            <li>Selecione um documento (DOCX, PDF, XLSX, CSV)</li>
            <li>Clique em <b>Processar</b> para extrair dados</li>
            <li>Visualize tabelas, listas e pares chave-valor</li>
        </ol>

        <h2>🌐 Aba Web</h2>
        <ol>
            <li>Digite a URL completa do site</li>
            <li>Clique em <b>Escanear</b></li>
            <li>Visualize formulários, tabelas e dados</li>
        </ol>

        <h2>🤖 Aba IA</h2>
        <ol>
            <li>Insira sua chave da API Gemini:
            <a href='https://aistudio.google.com/app/apikey'>https://aistudio.google.com/app/apikey</a></li>
            <li>Comandos sugeridos:</li>
            <ul>
                <li><i>Crie um contrato em PDF</i></li>
                <li><i>Gere uma planilha de orçamento</i></li>
                <li><i>Faça um relatório em DOCX</i></li>
            </ul>
        </ol>

        <h2>💡 Dicas Importantes</h2>
        <ul>
            <li>Para arquivos grandes, processe por partes</li>
            <li>Seja específico com a IA para melhores resultados</li>
            <li>Use exportação para compartilhar os resultados</li>
        </ul>
        """
        tutorial_text.setHtml(tutorial_html)
        tutorial_layout.addWidget(tutorial_text)
        tab_widget.addTab(tutorial_tab, "Tutorial")

        # === BOTÃO OK ===
        button_box = QDialogButtonBox(QDialogButtonBox.StandardButton.Ok)
        button_box.accepted.connect(dialog.accept)

        main_layout = QVBoxLayout(dialog)
        main_layout.addWidget(tab_widget)
        main_layout.addWidget(button_box)
        main_layout.setContentsMargins(10, 10, 10, 10)

        dialog.exec()

    def toggle_dark_mode(self):
        self.dark_mode = not self.dark_mode
        if self.dark_mode:
            self.apply_dark_theme()
        else:
            self.apply_light_theme()

    def apply_dark_theme(self):
        dark_palette = QPalette()
        
        # Configuração da paleta de cores
        dark_palette.setColor(QPalette.ColorRole.Window, QColor(53, 53, 53))
        dark_palette.setColor(QPalette.ColorRole.WindowText, Qt.GlobalColor.white)
        dark_palette.setColor(QPalette.ColorRole.Base, QColor(35, 35, 35))
        dark_palette.setColor(QPalette.ColorRole.AlternateBase, QColor(53, 53, 53))
        dark_palette.setColor(QPalette.ColorRole.ToolTipBase, Qt.GlobalColor.white)
        dark_palette.setColor(QPalette.ColorRole.ToolTipText, Qt.GlobalColor.white)
        dark_palette.setColor(QPalette.ColorRole.Text, Qt.GlobalColor.white)
        dark_palette.setColor(QPalette.ColorRole.Button, QColor(53, 53, 53))
        dark_palette.setColor(QPalette.ColorRole.ButtonText, Qt.GlobalColor.white)
        dark_palette.setColor(QPalette.ColorRole.BrightText, Qt.GlobalColor.red)
        dark_palette.setColor(QPalette.ColorRole.Link, QColor(42, 130, 218))
        dark_palette.setColor(QPalette.ColorRole.Highlight, QColor(42, 130, 218))
        dark_palette.setColor(QPalette.ColorRole.HighlightedText, Qt.GlobalColor.black)
        
        self.setPalette(dark_palette)
        
        self.setStyleSheet("""
            QTextEdit, QLineEdit, QComboBox {
                background-color: #252525;
                color: #ffffff;
                border: 2px solid #444;
                border-radius: 8px;
                padding: 12px;
                font-size: 14px;
                selection-background-color: #3a3a3a;
            }
            QPushButton {
                background-color: #0078d4;
                color: white;
                border-radius: 6px;
                padding: 10px 20px;
                font-weight: bold;
                min-width: 80px;
            }
            QPushButton:hover {
                background-color: #006cbd;
            }
            QTabWidget::pane {
                border: 1px solid #444;
                background: #353535;
            }
            QTabBar::tab {
                background: #353535;
                color: white;
                padding: 8px;
                border-top-left-radius: 4px;
                border-top-right-radius: 4px;
            }
            QTabBar::tab:selected {
                background: #252525;
                border-bottom-color: #0078d4;
            }
            QMenuBar {
                background-color: #353535;
                color: white;
            }
            QMenuBar::item {
                background-color: transparent;
                padding: 5px 10px;
            }
            QMenuBar::item:selected {
                background-color: #555;
            }
            QMenu {
                background-color: #353535;
                color: white;
                border: 1px solid #444;
            }
            QMenu::item:selected {
                background-color: #0078d4;
            }
        """)

    def apply_light_theme(self):
        self.setPalette(self.style().standardPalette())
        
        self.setStyleSheet("""
            QTextEdit, QLineEdit, QComboBox {
                background-color: #ffffff;
                color: #333333;
                border: 2px solid #0078d4;
                border-radius: 8px;
                padding: 12px;
                font-size: 14px;
                selection-background-color: #b8d8ff;
            }
            QPushButton {
                background-color: #0078d4;
                color: white;
                border-radius: 6px;
                padding: 10px 20px;
                font-weight: bold;
                min-width: 80px;
            }
            QPushButton:hover {
                background-color: #006cbd;
            }
            QTabWidget::pane {
                border: 1px solid #ddd;
                background: #f0f0f0;
            }
            QTabBar::tab {
                background: #f0f0f0;
                color: #333;
                padding: 8px;
                border-top-left-radius: 4px;
                border-top-right-radius: 4px;
            }
            QTabBar::tab:selected {
                background: #ffffff;
                border-bottom-color: #0078d4;
            }
            QMenuBar {
                background-color: #f0f0f0;
                color: #333;
            }
            QMenuBar::item:selected {
                background-color: #e0e0e0;
            }
            QMenu {
                background-color: #ffffff;
                color: #333;
                border: 1px solid #ddd;
            }
            QMenu::item:selected {
                background-color: #0078d4;
                color: white;
            }
        """)

    def init_file_tab(self, tab):
        layout = QVBoxLayout(tab)
        
        # File Selection
        file_layout = QHBoxLayout()
        self.file_path = QLineEdit()
        btn_browse = QPushButton("Selecionar Arquivo")
        btn_browse.clicked.connect(self.browse_file)
        
        file_layout.addWidget(self.file_path)
        file_layout.addWidget(btn_browse)
        
        # Preview
        self.preview = QTextEdit()
        
        # Processing
        btn_process = QPushButton("Processar")
        btn_process.clicked.connect(self.process_file)
        
        layout.addLayout(file_layout)
        layout.addWidget(self.preview)
        layout.addWidget(btn_process)

    def init_web_tab(self, tab):
        layout = QVBoxLayout(tab)
        
        # URL Input
        url_layout = QHBoxLayout()
        self.url_input = QLineEdit()
        btn_scrape = QPushButton("Escanear")
        btn_scrape.clicked.connect(self.scrape_web)
        
        url_layout.addWidget(QLabel("URL:"))
        url_layout.addWidget(self.url_input)
        url_layout.addWidget(btn_scrape)
        
        # Results
        self.web_results = QTextEdit()
        
        # Export Buttons
        export_layout = QHBoxLayout()
        btn_export_pdf = QPushButton("Exportar PDF")
        btn_export_excel = QPushButton("Exportar Excel")
        btn_export_docx = QPushButton("Exportar DOCX")
        
        btn_export_pdf.clicked.connect(lambda: self.export_web_results('pdf'))
        btn_export_excel.clicked.connect(lambda: self.export_web_results('excel'))
        btn_export_docx.clicked.connect(lambda: self.export_web_results('docx'))
        
        export_layout.addWidget(btn_export_pdf)
        export_layout.addWidget(btn_export_excel)
        export_layout.addWidget(btn_export_docx)
        
        layout.addLayout(url_layout)
        layout.addWidget(self.web_results)
        layout.addLayout(export_layout)

    def init_ai_tab(self, tab):
        layout = QVBoxLayout(tab)
        
        # API Key Section
        key_layout = QHBoxLayout()
        self.api_key_input = QLineEdit()
        self.api_key_input.setPlaceholderText("Cole sua API Key do Gemini aqui...")
        btn_save_key = QPushButton("🔑 Salvar")
        btn_save_key.clicked.connect(self.save_api_key)
        
        key_layout.addWidget(self.api_key_input)
        key_layout.addWidget(btn_save_key)
        
        # Chat history - Alterado para QTextBrowser
        self.chat_history = QTextBrowser()
        self.chat_history.setOpenExternalLinks(True)
        self.chat_history.setReadOnly(True)
        
        # User input
        input_layout = QHBoxLayout()
        self.ai_input_ia = QLineEdit()
        self.ai_input_ia.setPlaceholderText("Digite sua solicitação...")
        
        # Conecte o evento de pressionar Enter
        self.ai_input_ia.returnPressed.connect(self.handle_ai_input)
        
        btn_send = QPushButton("🚀 Enviar")
        btn_send.clicked.connect(self.handle_ai_input)
        
        input_layout.addWidget(self.ai_input_ia)
        input_layout.addWidget(btn_send)
        
        layout.addLayout(key_layout)
        layout.addWidget(self.chat_history)
        layout.addLayout(input_layout)

    def save_api_key(self):
        key = self.api_key_input.text().strip()
        if key:
            self.ai = AITerminal(api_key=key)
            QMessageBox.information(self, "Sucesso", "Chave configurada com sucesso!")
        else:
            QMessageBox.warning(self, "Aviso", "Insira uma API Key válida!")

    def browse_file(self):
        path, _ = QFileDialog.getOpenFileName(
            self, "Selecionar Arquivo", "", 
            "Documentos (*.docx *.pdf *.xlsx *.xls *.csv)"
        )
        if path:
            self.file_path.setText(path)
            self.preview.setText(f"Arquivo selecionado: {path}")

    def process_file(self):
        path = self.file_path.text()
        if not path:
            QMessageBox.warning(self, "Aviso", "Selecione um arquivo primeiro!")
            return
            
        try:
            result = process_file(path)
            self.preview.setText(result)
        except Exception as e:
            QMessageBox.critical(self, "Erro", f"Falha ao processar arquivo:\n{str(e)}")

    def scrape_web(self):
        url = self.url_input.text()
        if not url:
            QMessageBox.warning(self, "Aviso", "Digite uma URL válida!")
            return
            
        try:
            result = scrape_website(url)
            self.web_results.setText(result)
        except Exception as e:
            QMessageBox.critical(self, "Erro", f"Falha ao escanear website:\n{str(e)}")

    def handle_ai_input(self):
        # Obter texto e limpar imediatamente
        user_text = self.ai_input_ia.text().strip()
        self.ai_input_ia.clear()
        
        if not user_text:
            return
            
        self.chat_history.append(f"Você: {user_text}")
        
        # Foco de volta no campo de entrada
        self.ai_input_ia.setFocus()
        
        try:
            response, file_path = self.ai.generate_response(user_text)
            
            if file_path:
                self.current_file_path = file_path
                self.chat_history.append(
                    f"IA: {response}\n"
                    f"<a href='file:///{file_path}'>Abrir Documento</a>"
                )
            else:
                self.chat_history.append(f"IA: {response}")
            
            # Auto-scroll para a última mensagem
            scrollbar = self.chat_history.verticalScrollBar()
            scrollbar.setValue(scrollbar.maximum())
            
        except Exception as e:
            self.chat_history.append(f"IA: ❌ Erro ao processar solicitação: {str(e)}")
            QMessageBox.critical(self, "Erro", f"Falha na comunicação com a IA:\n{str(e)}")
    
    def export_web_results(self, format):
        content = self.web_results.toPlainText()
        if not content:
            QMessageBox.warning(self, "Aviso", "Nenhum resultado para exportar!")
            return
        
        filename, _ = QFileDialog.getSaveFileName(
            self, "Salvar Resultado", "", 
            f"{format.upper()} Files (*.{format})"
        )
        
        if not filename:
            return
        
        try:
            if format == 'pdf':
                from fpdf import FPDF
                pdf = FPDF()
                pdf.add_page()
                pdf.set_font("Arial", size=12)
                
                # Adiciona suporte a UTF-8
                pdf.set_doc_option('core_fonts_encoding', 'utf-8')
                
                # Quebra o texto em linhas de 200 caracteres
                lines = []
                current_line = ""
                for word in content.split():
                    if len(current_line + word) < 200:
                        current_line += word + " "
                    else:
                        lines.append(current_line)
                        current_line = word + " "
                lines.append(current_line)
                
                for line in lines:
                    pdf.multi_cell(0, 10, line)
                pdf.output(filename)
                
            elif format == 'excel':
                import pandas as pd
                data = []
                current_section = None
                
                # Processa o conteúdo de forma mais flexível
                for line in content.split('\n'):
                    line = line.strip()
                    if not line:
                        continue
                    
                    if line.startswith('==='):
                        current_section = line.replace('=', '').strip()
                    elif ':' in line:
                        key, value = line.split(':', 1)
                        data.append({
                            'Seção': current_section or 'Geral',
                            'Chave': key.strip(),
                            'Valor': value.strip()
                        })
                
                df = pd.DataFrame(data)
                df.to_excel(filename, index=False)
                
            elif format == 'docx':
                from docx import Document
                doc = Document()
                doc.add_heading('Resultado da Web Scraping', 0)
                
                # Adiciona parágrafos preservando quebras de linha
                for line in content.split('\n'):
                    if line.strip():
                        doc.add_paragraph(line.strip())
                
                doc.save(filename)
                
            QMessageBox.information(self, "Sucesso", f"Arquivo salvo em:\n{filename}")
            
        except Exception as e:
            QMessageBox.critical(self, "Erro", f"Falha ao exportar:\n{str(e)}")

if __name__ == "__main__":
    from PyQt6.QtWidgets import QApplication
    app = QApplication([])
    
    # Configuração para alta DPI (melhor visualização em monitores 4K)
    app.setAttribute(Qt.ApplicationAttribute.AA_EnableHighDpiScaling)
    app.setAttribute(Qt.ApplicationAttribute.AA_UseHighDpiPixmaps)
    
    window = DataScan()
    window.show()
    app.exec()