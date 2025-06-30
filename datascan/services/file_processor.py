from docx import Document
import pandas as pd
import PyPDF2
import os
import re
from collections import defaultdict

def process_file(file_path):
    try:
        if file_path.endswith('.docx'):
            return process_docx(file_path)
        elif file_path.endswith('.pdf'):
            return process_pdf(file_path)
        elif file_path.endswith(('.xlsx', '.xls', '.csv')):
            return process_excel(file_path)
        else:
            return "Formato não suportado"
    except Exception as e:
        return f"Erro: {str(e)}"

def process_docx(path):
    doc = Document(path)
    structured_data = defaultdict(list)
    
    # 1. Extrair parágrafos com padrões chave:valor
    for para in doc.paragraphs:
        text = para.text.strip()
        if re.match(r'^[^:]+:[^:]+$', text):  # Padrão "chave: valor"
            key, value = text.split(':', 1)
            structured_data['key_value'].append((key.strip(), value.strip()))
    
    # 2. Extrair tabelas
    for table in doc.tables:
        table_data = []
        headers = [cell.text.strip() for cell in table.rows[0].cells]
        
        for row in table.rows[1:]:
            row_data = {}
            for idx, cell in enumerate(row.cells):
                if idx < len(headers):
                    row_data[headers[idx]] = cell.text.strip()
            table_data.append(row_data)
        
        structured_data['tables'].append({
            'headers': headers,
            'rows': table_data
        })
    
    # 3. Extrair listas
    list_items = []
    for para in doc.paragraphs:
        if para.style.name.startswith('List'):
            list_items.append(para.text.strip())
    if list_items:
        structured_data['lists'].append(list_items)
    
    return format_structured_data(structured_data, os.path.basename(path))

def process_pdf(path):
    structured_data = defaultdict(list)
    with open(path, 'rb') as file:
        reader = PyPDF2.PdfReader(file)
        
        for page_num in range(len(reader.pages)):
            page = reader.pages[page_num]
            text = page.extract_text()
            
            # 1. Extrair chave:valor
            for line in text.split('\n'):
                if re.match(r'^[^:]+:[^:]+$', line.strip()):
                    key, value = line.split(':', 1)
                    structured_data['key_value'].append((key.strip(), value.strip()))
            
            # 2. Tentar identificar tabelas (abordagem heurística)
            table_lines = []
            for line in text.split('\n'):
                if re.match(r'(\s{2,}.+)+', line):  # Linhas com múltiplos espaços
                    table_lines.append(re.split(r'\s{2,}', line.strip()))
            
            if len(table_lines) > 1:
                headers = table_lines[0]
                table_data = []
                for row in table_lines[1:]:
                    if len(row) == len(headers):
                        table_data.append(dict(zip(headers, row)))
                
                structured_data['tables'].append({
                    'headers': headers,
                    'rows': table_data
                })
    
    return format_structured_data(structured_data, os.path.basename(path))

def process_excel(path):
    if path.endswith(('.xlsx', '.xls')):
        df = pd.read_excel(path)
    else:
        df = pd.read_csv(path)
    
    structured_data = defaultdict(list)
    
    # 1. Extrair como tabela principal
    structured_data['tables'].append({
        'headers': df.columns.tolist(),
        'rows': df.head(20).to_dict('records')  # Limitar a 20 linhas
    })
    
    # 2. Procurar colunas com padrões chave:valor
    for col in df.columns:
        if any(df[col].astype(str).str.contains(r'^[^:]+:[^:]+$', na=False)):
            key_value_pairs = []
            for val in df[col]:
                if isinstance(val, str) and ':' in val:
                    key, value = val.split(':', 1)
                    key_value_pairs.append((key.strip(), value.strip()))
            structured_data['key_value'].extend(key_value_pairs)
    
    return format_structured_data(structured_data, os.path.basename(path))

def format_structured_data(data, filename):
    """Formata os dados estruturados similar ao web_scraper"""
    result = f"=== ESTRUTURA DO ARQUIVO: {filename} ===\n\n"
    
    if data['key_value']:
        result += "=== PARES CHAVE-VALOR ===\n"
        for key, value in data['key_value']:
            result += f"{key}: {value}\n"
        result += "\n"
    
    if data['tables']:
        result += "=== TABELAS ===\n"
        for table_idx, table in enumerate(data['tables'], 1):
            result += f"\nTabela {table_idx}:\n"
            result += " | ".join(table['headers']) + "\n"
            result += "-" * (sum(len(h) for h in table['headers']) + 3*len(table['headers'])) + "\n"
            
            for row in table['rows'][:5]:  # Limitar a 5 linhas
                values = [str(row.get(h, '')) for h in table['headers']]
                result += " | ".join(values) + "\n"
            result += f"\nTotal de linhas: {len(table['rows'])}\n\n"
    
    if data['lists']:
        result += "=== LISTAS ===\n"
        for list_idx, items in enumerate(data['lists'], 1):
            result += f"\nLista {list_idx} ({len(items)} itens):\n"
            for item in items[:10]:  # Limitar a 10 itens
                result += f"  - {item}\n"
    
    # Salvar relatório
    report_path = f"structured_report_{os.path.splitext(filename)[0]}.txt"
    with open(report_path, 'w', encoding='utf-8') as f:
        f.write(result)
    
    return f"Relatório gerado: {report_path}\n\n{result}"