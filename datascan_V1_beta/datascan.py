import os
import json
import torch
import gc
import tkinter as tk
import zipfile
from tkinter import filedialog, messagebox, simpledialog, ttk
from PyPDF2 import PdfReader
from docx import Document
from openpyxl import load_workbook
from openpyxl.styles import Font
from transformers import AutoTokenizer, AutoModelForCausalLM, AutoModelForQuestionAnswering
from bs4 import BeautifulSoup
import validators
import requests
import threading
from queue import Queue
from urllib.parse import urlparse
import time
import pandas as pd
from datetime import datetime
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter
from reportlab.platypus import Table, TableStyle
from reportlab.lib import colors
import matplotlib.pyplot as plt
from matplotlib.backends.backend_tkagg import FigureCanvasTkAgg
# Adicione no topo do arquivo com os outros imports
from openpyxl import Workbook
import re
import psutil

import sys


# Configurações de baixo consumo
torch.backends.cudnn.benchmark = False
torch.set_flush_denormal(True)

# Liberar memória antes de verificar
torch.cuda.empty_cache() if torch.cuda.is_available() else None
gc.collect()


def check_system_resources():
    mem = psutil.virtual_memory()
    available_gb = mem.available / (1024 ** 3)
    
    # Novo limite mínimo de 0.5GB
    if available_gb < 0.5:
        raise MemoryError(
            f"Memória crítica! Disponível: {available_gb:.2f} GB. "
            "Feche todos os outros programas e reinicie o computador."
        )
    elif available_gb < 1.5:
        print(f"[AVISO] Memória baixa: {available_gb:.2f} GB. Performance pode ser afetada.")

class Config:
    MODEL_NAME = "pierreguillou/gpt2-small-portuguese "
    PAD_TOKEN = "<|pad|>"
    MAX_TOKENS = 64
    MAX_CONTEXT_LENGTH = 2
    USE_QUANTIZATION = False
    DEVICE = "cpu"
    TEMP_DIR = "temp_processing"
    BACKUP_DIR = "backups"
    
    # Novo: Parâmetros de geração melhorados
    GENERATION_CONFIG = {
        "max_new_tokens": 150,  # Limita o tamanho da resposta
        "temperature": 0.3,    # Controla a criatividade (valores mais baixos = mais conservador)
        "top_p": 0.9,          # Filtra as opções de geração
        "top_k": 50,           # Limita o número de opções de geração
        "repetition_penalty": 2.0,  # Evita repetições
        "no_repeat_ngram_size": 3,
        "do_sample": True,     # Ativa a amostragem
        "num_beams": 1,        # Número de beams (1 para geração simples)
        "pad_token_id": 50256, # Token de padding
    }
    
    MODEL_LOAD_SETTINGS = {
        "torch_dtype": torch.float32,
        "low_cpu_mem_usage": True,
        "device_map": "cpu",
        "max_memory": {"cpu": "0.5GB"}
    }
    
     # Nova paleta de cores profissional
    PRIMARY_COLOR = "#2A3F5F"      # Azul escuro
    SECONDARY_COLOR = "#4A6572"    # Cinza azulado
    ACCENT_COLOR = "#F9AA33"       # Laranja dourado
    BACKGROUND_COLOR = "#E0E0E0"   # Cinza claro
    TEXT_COLOR = "#FFFFFF"         # Branco
    CHAT_BG = "#37474F"            # Cinza escuro
    HOVER_COLOR = "#566573"        # Efeito hover
    
# Garantir criação dos diretórios necessários
if not os.path.exists(Config.TEMP_DIR):
    os.makedirs(Config.TEMP_DIR)
if not os.path.exists(Config.BACKUP_DIR):
    os.makedirs(Config.BACKUP_DIR)
    

class DocumentProcessor:
    def __init__(self, master=None):
        # Corrigir ordem de inicialização
        self.tokenizer = AutoTokenizer.from_pretrained(
            "pierreguillou/gpt2-small-portuguese",  # Modelo generativo em português
            use_fast=True
        )
        
        # Adicionar token de padding APÓS a inicialização
        self.tokenizer.add_special_tokens({'pad_token': Config.PAD_TOKEN})
        
        self.model = AutoModelForCausalLM.from_pretrained(
            "pierreguillou/gpt2-small-portuguese",
            **Config.MODEL_LOAD_SETTINGS
        )
        
        # Redimensionar embeddings DEPOIS de carregar o modelo
        self.model.resize_token_embeddings(len(self.tokenizer))
        # Otimizações adicionais
        self.model.eval()
        self.model = torch.quantization.quantize_dynamic(
            self.model,
            {torch.nn.Linear},
            dtype=torch.qint8
        )
        
        self.chat_history = []
        self.template_generator = TemplateGenerator()
        self.template_generator = TemplateGenerator(master)
        # Gerenciamento de memória
        #torch.set_num_threads(1)  # Limita uso de CPU
        
        
    def load_model(self):
        if self.model is None:
            self.tokenizer = AutoTokenizer.from_pretrained(Config.MODEL_NAME)
            self.model = AutoModelForCausalLM.from_pretrained(  # ← Alterado
                Config.MODEL_NAME,
                torch_dtype=torch.float32,
                low_cpu_mem_usage=False
            )
            self.model.eval()
        
    def _classify_intent(self, prompt: str) -> str:
        prompt_lower = prompt.lower().strip()
        
        # Detecção direta de tipos de documento
        if prompt_lower in {'doc', 'docx', 'xls', 'xlsx'}:
            return 'template'
        
        intents = {
            'template': ['template', 'modelo', 'criar', 'gerar'],
            'process': ['processar', 'arquivo', 'documento'],
            'analyze': ['analisar', 'site', 'web'],
            'help': ['ajuda', 'comandos', '?']
        }
        
        for intent, keywords in intents.items():
            if any(kw in prompt_lower for kw in keywords):
                return intent
                
        return 'general'

    def _handle_multi_step(self, prompt: str) -> str:
        """Gerencia fluxos de diálogo com múltiplas etapas."""
        if self.dialog_context['last_action'] == 'template_creation':
            return self._handle_template_flow(prompt)
            
        elif self.dialog_context['last_action'] == 'file_processing':
            return self._handle_file_flow(prompt)
            
        elif self.dialog_context['last_action'] == 'website_analysis':
            return self._handle_website_flow(prompt)
            
        return "Por favor, continue sua solicitação."

    def _handle_template_flow(self, prompt: str) -> str:
        """Gerencia o fluxo de criação de templates."""
        if self.dialog_context['step'] == 0:
            self.dialog_context['step'] = 1
            return (
                "⚠️ Por favor, forneça ESTAS informações:\n"
                "1. Tipo (DOC ou XLSX)\n"
                "2. Campos separados por vírgula\n"
                "Exemplo: 'DOC, Nome, Data, Valor'"
            )
            
        elif self.dialog_context['step'] == 1:
            try:
                # Validação reforçada
                if not re.match(r'^(docx?|xlsx),\s*.+', prompt, re.IGNORECASE):
                    raise ValueError(
                        "Formato inválido! Use: TIPO, Campo1, Campo2...\n"
                        "Ex: 'DOC, Cliente, Valor' ou 'XLSX, Item, Quantidade'"
                    )
                
                parts = [p.strip() for p in prompt.split(',')]
                template_type = parts[0].lower().replace("doc", "docx")  # Normaliza para docx
                fields = [f for f in parts[1:] if f]
                
                # Validação de tipos
                if template_type not in ['docx', 'xlsx']:
                    raise ValueError("Tipo inválido. Use DOC/DOCX ou XLSX")
                    
                if not fields:
                    raise ValueError("Pelo menos um campo é necessário")
                    
                # Geração do template
                output_path = self.generate_template(template_type, fields)
                
                self.dialog_context = {'last_action': None, 'pending_data': None, 'step': 0}
                return (
                    f"✅ Template {template_type.upper()} criado!\n"
                    f"📂 Local: {output_path}\n"
                    f"🗂️ Campos: {', '.join(fields)}"
                )
                
            except Exception as e:
                self.dialog_context = {'last_action': None, 'pending_data': None, 'step': 0}
                return f"❌ Falha na criação: {str(e)}"

    def _handle_file_flow(self, prompt: str) -> str:
        """Gerencia o fluxo de processamento de arquivos."""
        if self.dialog_context['step'] == 0:
            self.dialog_context['step'] = 1
            return "Por favor, selecione o arquivo que deseja processar."
            
        elif self.dialog_context['step'] == 1:
            try:
                content = self.extract_from_file(prompt)
                self.document_content = content
                self.dialog_context['step'] = 2
                return (
                    "Arquivo processado com sucesso!\n"
                    "Agora, selecione o template para gerar a saída."
                )
            except Exception as e:
                self.dialog_context = {'last_action': None, 'pending_data': None, 'step': 0}
                return f"❌ Erro no processamento: {str(e)}"

    def quantize_model(self, model):
        return torch.quantization.quantize_dynamic(
            model,
            {torch.nn.Linear},
            dtype=torch.qint8
        )

    def extract_from_file(self, file_path: str) -> str:
        try:
            # Verificação de tamanho do arquivo
            max_size = 50 * 1024 * 1024  # 50MB
            file_size = os.path.getsize(file_path)
            if file_size > max_size:
                raise ValueError(f"Arquivo muito grande ({file_size/1024/1024:.2f}MB > 50MB)")
                
            chunk_size = 1024
            content = []
            ext = os.path.splitext(file_path)[1].lower()

            # Leitura segura para cada formato
            if ext == '.pdf':
                reader = PdfReader(file_path)
                for page in reader.pages[:10]:
                    text = page.extract_text() or ""
                    # Corrigido: Fechamento correto dos colchetes e parênteses
                    content.extend([text[i:i+chunk_size] for i in range(0, len(text), chunk_size)])

            elif ext == '.docx':
                doc = Document(file_path)
                for para in doc.paragraphs[:100]:
                    if para.text.strip():
                        content.append(para.text[:chunk_size].replace('\n', ' '))

            elif ext in ('.xlsx', '.xls'):
                wb = load_workbook(file_path, read_only=True)
                for sheet in wb:
                    for row in sheet.iter_rows(values_only=True)[:1000]:
                        row_content = [str(cell)[:50] for cell in row if cell is not None]
                        content.append(" | ".join(row_content))
                    break  # Processa apenas a primeira planilha

            elif ext == '.txt':
                with open(file_path, 'r', encoding='utf-8', errors='replace') as f:
                    while True:
                        chunk = f.read(chunk_size)
                        if not chunk:
                            break
                        content.append(chunk.replace('\n', ' '))
                        if len(content) >= 1000:  # Limite de chunks
                            break
            else:
                raise ValueError(f"Formato não suportado: {ext}")

            full_content = " ".join(content)
            return full_content[:Config.MAX_TOKENS * 4].strip()

        except Exception as e:
            raise Exception(f"Erro na extração: {str(e)}")
    def export_to_pdf(self, data: dict, filename: str):
        """Exporta dados para PDF com formatação profissional"""
        try:
            c = canvas.Canvas(filename, pagesize=letter)
            width, height = letter
            
            # Cabeçalho
            c.setFillColorRGB(0, 1, 0)  # Verde
            c.setFont("Helvetica-Bold", 16)
            c.drawString(100, height - 100, "Relatório Gerado por DataScan")
            c.line(100, height - 110, width - 100, height - 110)
            
            # Conteúdo
            y_position = height - 150
            c.setFont("Helvetica", 12)
            c.setFillColorRGB(1, 1, 1)  # Branco
            
            # Tabela de dados
            if 'data' in data and len(data['data']) > 0:
                table_data = [list(data['data'][0].keys())]  # Cabeçalhos
                for row in data['data']:
                    table_data.append([str(v) for v in row.values()])
                
                t = Table(table_data)
                t.setStyle(TableStyle([
                    ('BACKGROUND', (0,0), (-1,0), colors.green),
                    ('TEXTCOLOR', (0,0), (-1,0), colors.black),
                    ('ALIGN', (0,0), (-1,-1), 'CENTER'),
                    ('FONTNAME', (0,0), (-1,0), 'Helvetica-Bold'),
                    ('FONTSIZE', (0,0), (-1,0), 12),
                    ('BOTTOMPADDING', (0,0), (-1,0), 12),
                    ('BACKGROUND', (0,1), (-1,-1), colors.black),
                    ('TEXTCOLOR', (0,1), (-1,-1), colors.green),
                    ('GRID', (0,0), (-1,-1), 1, colors.green)
                ]))
                
                t.wrapOn(c, width-200, height)
                t.drawOn(c, 100, y_position - len(table_data)*20)
                y_position -= len(table_data)*20 + 50

            # Análise da IA
            if 'analysis' in data:
                c.setFont("Helvetica", 10)
                c.drawString(100, y_position, "Análise da IA:")
                text_object = c.beginText(100, y_position - 20)
                text_object.setFont("Helvetica", 10)
                text_object.setFillColor(colors.green)
                
                for line in data['analysis'].split('\n'):
                    text_object.textLine(line)
                    y_position -= 12
                    if y_position < 100:
                        c.showPage()
                        y_position = height - 100
                        text_object = c.beginText(100, y_position)
                        
                c.drawText(text_object)

            c.save()
            return True
        except Exception as e:
            raise Exception(f"Erro ao gerar PDF: {str(e)}")

    def scrape_website(self, url: str) -> dict:
        try:
            # Configuração de sessão com retries e headers
            session = requests.Session()
            session.mount('https://', requests.adapters.HTTPAdapter(
                max_retries=3,
                pool_maxsize=10
            ))
            headers = {
                'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64)',
                'Accept-Language': 'pt-BR,pt;q=0.9',
                'Accept-Encoding': 'gzip, deflate'
            }

            # Validação e normalização da URL
            url = url.strip().replace(" ", "")
            if not url:
                raise ValueError("URL não pode estar vazia")
            
            if not url.startswith(('http://', 'https://')):
                url = f'https://{url}'
                
            if not validators.url(url):
                raise ValueError(f"URL inválida: {url}")
                
            parsed_url = urlparse(url)
            if not validators.domain(parsed_url.hostname):
                raise ValueError(f"Domínio inválido: {parsed_url.hostname}")

            # Requisição com tratamento de segurança
            response = session.get(
                url,
                headers=headers,
                timeout=20,
                verify=True  # Alterar para False se necessário em ambientes controlados
            )
            response.raise_for_status()

            # Limitação de tamanho do conteúdo
            max_content_size = 5000000  # 5MB
            if len(response.content) > max_content_size:
                raise ValueError("Conteúdo do site excede o tamanho máximo permitido (5MB)")

            soup = BeautifulSoup(response.text, 'html.parser')
            raw_text = soup.get_text(separator=' ', strip=True)[:4000]

            # Processamento com IA com fallback
            try:
                ai_prompt = ("Extraia informações estruturadas deste site...")
                ai_response = self.generate_ai_response(ai_prompt)
                structured_data = json.loads(ai_response)
            except Exception as ai_error:
                structured_data = {"error": str(ai_error), "content": raw_text[:2000]}

            return {
                "title": soup.title.string if soup.title else "Sem título",
                "headers": [header.text.strip() for header in soup.find_all(['h1', 'h2', 'h3'])],
                "links": [link.get('href') for link in soup.find_all('a') if link.get('href')],
                "content": raw_text,
                "structured_data": structured_data,
                "ai_analysis": ai_response
            }
            
        except Exception as e:
            raise Exception(f"Erro no scraping: {str(e)}")
        
    def generate_ai_response(self, prompt: str) -> str:
        try:
            # Verifica fluxo de template ativo
            if hasattr(self, 'template_generator') and self.template_generator.context['step'] > 0:
                response = self.template_generator.handle_input(prompt)
                
                # Adiciona sugestões inteligentes
                if self.template_generator.context['step'] == 1:  # Durante entrada de campos
                    enhanced = self._enhance_with_ai_suggestions(prompt)
                    if enhanced:
                        response += f"\n\n💡 Sugestão da IA: {enhanced}"
                return response
            
            # Processamento inteligente do prompt
            intent = self._classify_intent(prompt)
            
            if intent == 'template':
                if not hasattr(self, 'template_generator'):
                    self.template_generator = TemplateGenerator()
                
                # Geração contextual
                if "relatório" in prompt.lower():
                    suggestion = "Considere incluir: Introdução*, Métodos*, Resultados*"
                elif "contrato" in prompt.lower():
                    suggestion = "Campos recomendados: Partes*, Objeto*, Vigência*"
                else:
                    suggestion = ""
                
                if ',' in prompt:
                    return f"{self.template_generator.handle_input(prompt)}\n{suggestion}"
                else:
                    return self.template_generator.handle_input(prompt)
            
            return self._handle_general_query(prompt)
            
        except Exception as e:
            return f"❌ Erro: {str(e)}"
        
    def _handle_general_query(self, prompt: str) -> str:
        """Processa consultas gerais e saudações"""
        greetings = {'olá', 'oi', 'bom dia', 'boa tarde', 'boa noite'}
        if any(g in prompt.lower() for g in greetings):
            return self._handle_greeting(prompt)
        return "Desculpe, não entendi. Digite 'ajuda' para ver opções disponíveis."
        
    def _enhance_with_ai_suggestions(self, prompt: str) -> str:
        """Gera sugestões de campos baseadas no contexto"""
        try:
            inputs = self.tokenizer(
                f"Sugira campos adicionais para: {prompt}",
                return_tensors="pt",
                truncation=True
            )
            outputs = self.model.generate(
                **inputs,
                max_new_tokens=50,
                temperature=0.7
            )
            return self.tokenizer.decode(outputs[0], skip_special_tokens=True)
        except:
            return ""

            
    # Adicionar após a geração da resposta:
    def _filter_response(self, response: str) -> str:
        """Filtra conteúdo técnico e lixo gerado"""
        filters = [
            r'http[s]?://(?:[a-zA-Z]|[0-9]|[$-_@.&+]|[!*\\(\\),]|(?:%[0-9a-fA-F][0-9a-fA-F]))+',
            r'\b[A-Z]{5,}\b',  # Palavras totalmente em maiúsculo
            r'[^\w\s,.!?À-ÿ]',  # Caracteres especiais
            r'\b\w{20,}\b',  # Palavras muito longas
            r'([A-Za-z])\1{3,}'  # Letras repetidas
        ]
        
        for pattern in filters:
            response = re.sub(pattern, '', response)
        
        return response.strip()[:500]  # Limite de caracteres


    def generate_template(self, template_type: str, fields: list) -> str:
        try:
            timestamp = int(time.time())
            
            if template_type == 'docx':
                from docx.shared import Pt, RGBColor, Inches
                from docx.enum.text import WD_ALIGN_PARAGRAPH
                from docx.enum.style import WD_STYLE_TYPE
                
                doc = Document()
                
                # ========== ESTILOS PERSONALIZADOS ==========
                styles = doc.styles
                
                # Estilo Título
                title_style = styles.add_style('CustomTitle', WD_STYLE_TYPE.PARAGRAPH)
                title_font = title_style.font
                title_font.name = 'Calibri Light'
                title_font.size = Pt(18)
                title_font.color.rgb = RGBColor(0x2A, 0x3F, 0x5F)  # Azul escuro
                title_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
                title_style.paragraph_format.space_after = Pt(24)
                
                # Estilo Cabeçalho
                header_style = styles.add_style('CustomHeader', WD_STYLE_TYPE.PARAGRAPH)
                header_font = header_style.font
                header_font.name = 'Calibri'
                header_font.size = Pt(10)
                header_font.color.rgb = RGBColor(0x7F, 0x7F, 0x7F)  # Cinza
                header_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    
                # Estilo Rodapé
                footer_style = styles.add_style('CustomFooter', WD_STYLE_TYPE.PARAGRAPH)
                footer_font = footer_style.font
                footer_font.name = 'Calibri'
                footer_font.size = Pt(9)
                footer_font.color.rgb = RGBColor(0x7F, 0x7F, 0x7F)  # Cinza
                footer_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
                footer_style.paragraph_format.space_before = Pt(6)
                
                # ========== CABEÇALHO DO DOCUMENTO ==========
                header = doc.sections[0].header
                header_para = header.add_paragraph()
                header_para.style = 'CustomHeader'
                header_run = header_para.add_run("Template Gerado - V1.0")
                
                # ========== TÍTULO PRINCIPAL ==========
                title = doc.add_paragraph('FORMULÁRIO PADRONIZADO', style='CustomTitle')
                
                # ========== METADADOS ==========
                meta_table = doc.add_table(rows=2, cols=2)
                meta_table.style = 'LightGrid-Accent1'
                meta_table.autofit = True
                meta_table.columns[0].width = Inches(1.5)
                
                # Linha 1 - Data
                meta_table.cell(0, 0).text = "Data:"
                meta_table.cell(0, 1).text = datetime.now().strftime("%d/%m/%Y")
                
                # Linha 2 - Versão
                meta_table.cell(1, 0).text = "Versão:"
                meta_table.cell(1, 1).text = "1.0"
                
                doc.add_paragraph()  # Espaçamento
                
                # ========== SEÇÃO DE CAMPOS ==========
                for field in fields:  # LOOP CRÍTICO QUE ESTAVA FALTANDO
                    # Dividir campo e formatação
                    if '*' in field:
                        parts = field.split('*', 1)  # Divide apenas no primeiro *
                        field_name = parts[0].strip()
                        fmt = parts[1].strip().lower()
                    else:
                        field_name = field.strip()
                        fmt = None

                    # Aplicar formatação
                    if fmt == 'h1':
                        doc.add_heading(field_name, level=1)
                    elif fmt == 'h2':
                        doc.add_heading(field_name, level=2)
                    elif fmt == 'blt':
                        p = doc.add_paragraph(style='List Bullet')
                        p.add_run(f"{field_name}: ").bold = True
                        p.add_run("____________________")
                    elif fmt == 'num':
                        p = doc.add_paragraph(style='List Number')
                        p.add_run(f"{field_name}: ").bold = True
                        p.add_run("____________________")
                    elif fmt == 'qt':
                        p = doc.add_paragraph(style='IntenseQuote')
                        p.add_run(f"{field_name}: ").italic = True
                        p.add_run("____________________")
                    elif fmt == 'sig':
                        p = doc.add_paragraph()
                        p.add_run(f"{field_name}:").bold = True
                        p.add_run("\n\n_____________________________\n")
                        p.add_run("(Assinatura)").italic = True
                    elif fmt == 'date':
                        p = doc.add_paragraph()
                        p.add_run(f"{field_name}:").bold = True
                        p.add_run("\n" + datetime.now().strftime("%d/%m/%Y"))
                    else:  # Campo sem formatação especial
                        p = doc.add_paragraph()
                        p.add_run(f"{field_name}:").bold = True
                        p.add_run(" ____________________")

                output_path = os.path.join(Config.TEMP_DIR, f"template_{timestamp}.docx")
                doc.save(output_path)
                
            elif template_type == 'xlsx':
                from openpyxl.styles import Font, PatternFill, Border, Side, Alignment
                from openpyxl.utils import get_column_letter
                
                wb = Workbook()
                ws = wb.active
                ws.title = "Dados"
                
                # ========== ESTILOS ==========
                header_font = Font(name='Calibri', bold=True, size=12, color='FFFFFF')
                header_fill = PatternFill(start_color='2A3F5F', end_color='2A3F5F', fill_type='solid')
                cell_font = Font(name='Calibri', size=11)
                border = Border(
                    left=Side(style='thin'),
                    right=Side(style='thin'),
                    top=Side(style='thin'),
                    bottom=Side(style='thin')
                )
                center_alignment = Alignment(horizontal='center', vertical='center')
                
                # ========== CABEÇALHO ==========
                # Título
                ws.merge_cells('A1:D1')
                title_cell = ws['A1']
                title_cell.value = "PLANILHA PADRONIZADA"
                title_cell.font = Font(name='Calibri Light', size=14, bold=True, color='2A3F5F')
                title_cell.alignment = center_alignment
                
                # Metadados
                ws['A2'] = "Gerado em:"
                ws['B2'] = datetime.now().strftime("%d/%m/%Y %H:%M")
                ws['A3'] = "Versão:"
                ws['B3'] = "1.0"
                
                # ========== CAMPOS ==========
                # Cabeçalhos dos campos
                for col, field in enumerate(fields, 1):
                    cell = ws.cell(row=5, column=col, value=field)
                    cell.font = header_font
                    cell.fill = header_fill
                    cell.border = border
                    cell.alignment = center_alignment
                    
                    # Ajuste automático de largura
                    column_letter = get_column_letter(col)
                    ws.column_dimensions[column_letter].width = max(len(field) + 4, 15)
                
                # Área de dados com formatação condicional
                for row in range(6, 16):  # 10 linhas de dados
                    for col, field in enumerate(fields, 1):
                        cell = ws.cell(row=row, column=col)
                        cell.value = f'=IF(ISBLANK({get_column_letter(col)}5),"",{get_column_letter(col)}5)'
                        cell.font = cell_font
                        cell.border = border
                        
                        # Formatação condicional para campos obrigatórios
                        if col == 1:  # Primeiro campo como obrigatório
                            cell.fill = PatternFill(start_color='FFF2CC', end_color='FFF2CC', fill_type='solid')
                
                # ========== CONFIGURAÇÕES ADICIONAIS ==========
                ws.freeze_panes = 'A6'  # Congela cabeçalhos
                ws.print_options.horizontalCentered = True
                ws.print_options.verticalCentered = False
                
                output_path = os.path.join(Config.TEMP_DIR, f"template_{timestamp}.xlsx")
                wb.save(output_path)
                
            return output_path

        except Exception as e:
            raise Exception(f"Erro na geração: {str(e)}")

    def fill_template(self, template_path: str, output_path: str, data: dict) -> bool:
        try:
            # Combina dados tradicionais e estruturados
            full_data = {
                **data.get('structured_data', {}),
                **{f"raw_{k}": v for k, v in data.items()}
            }
            
            ext = os.path.splitext(template_path)[1].lower()
            if ext == '.docx':
                self.fill_docx_template(template_path, output_path, full_data)
            elif ext in ('.xlsx', '.xls'):
                self.fill_excel_template(template_path, full_data).save(output_path)
            elif ext == '.pdf':
                raise NotImplementedError("Preenchimento de PDF não implementado")
            return True
        except Exception as e:
            raise Exception(f"Erro no template: {str(e)}")

    def fill_docx_template(self, template_path: str, output_path: str, data: dict):
        doc = Document(template_path)
        for para in doc.paragraphs:
            if '{{' in para.text:
                for key, value in data.items():
                    if isinstance(value, list):
                        value = ', '.join(value)
                    para.text = para.text.replace(f'{{{{{key}}}}}', str(value))
        doc.save(output_path)

    def fill_excel_template(self, template_path: str, data: dict):
        wb = load_workbook(template_path)
        for sheet in wb:
            for row in sheet.iter_rows():
                for cell in row:
                    if cell.value and '{{' in str(cell.value):
                        for key, value in data.items():
                            if isinstance(value, list):
                                value = ', '.join(value)
                            cell.value = str(cell.value).replace(f'{{{{{key}}}}}', str(value))
        return wb
    
    def _handle_greeting(self, prompt: str) -> str:
        greetings = {
            'olá': "👋 Olá! Vamos criar um template profissional?",
            'oi': "👋 Oi! Pronto para criar templates incríveis?",
            'bom dia': "🌞 Bom dia! Vamos criar alguns templates?"
        }
        for g in greetings:
            if g in prompt.lower():
                return (
                    f"{greetings[g]}\n\n"
                    "📝 **Como criar um template**:\n"
                    "1. Digite o tipo (docx ou xlsx)\n"
                    "2. Liste os campos separados por vírgula\n"
                    "3. Adicione formatação com * (opcional)\n\n"
                    "💡 **Exemplos**:\n"
                    "- Para Word: 'docx, Título*h1, Nome, Data*date'\n"
                    "- Para Excel: 'xlsx, Produto, Quantidade, Preço*curr'\n\n"
                    "Digite o tipo de template que deseja criar:"
                )
        return ""

    
class TemplateGenerator:
    def __init__(self, master=None):
        self.master = master
        self.context = {'step': 0, 'template_type': None, 'fields': []}
        
        # Mapeamento de estilos abreviados
        self.style_shortcuts = {
            'h1': 'heading1',
            'h2': 'heading2',
            'blt': 'bullet',
            'qt': 'quote',
            'tit': 'heading1',
            'sub': 'heading2',
            'lst': 'bullet'
        }

        self.reset_context()
    
    def get_format_help(self, template_type: str) -> str:
        help_texts = {
            'docx': (
                "✏️ **FORMATAÇÃO PARA DOCUMENTOS**\n"
                "────────────────────────────────\n"
                "Adicione * após o campo com:\n\n"
                "• *h1 - Título principal\n"
                "• *h2 - Subtítulo\n"
                "• *blt - Lista com marcadores\n"
                "• *num - Lista numerada\n"
                "• *qt - Citação\n"
                "• *sig - Campo de assinatura\n"
                "• *date - Data automática\n"
                "• *table - Tabela pré-formatada\n\n"
                "💡 **Exemplo**: 'Título*h1, Itens*blt, Data*date, Assinatura*sig'"
            ),
            'xlsx': (
                "📊 **FORMATAÇÃO PARA PLANILHA EXCEL**\n"
                "───────────────────────────────────────\n"
                "Adicione * após o campo:\n\n"
                "• *curr - Formato monetário\n"
                "• *date - Formato de data\n\n"
                "💡 Exemplo: 'Produto, Quantidade, Preço*curr'\n"
                "Digite os campos separados por vírgulas:"
            )
        }
        return help_texts.get(template_type, "Selecione um tipo válido (docx/xlsx)")
        
       # return help_texts.get(template_type, "Selecione um tipo válido (docx/xlsx)")

    def handle_input(self, user_input: str) -> str:
        user_input = user_input.lower().strip()
        
        # Comandos de cancelamento em qualquer etapa
        if user_input in ['sair', 'cancelar', 'voltar']:
            self.reset_context()
            return "❌ Operação cancelada. O que deseja fazer agora?"
        
        # Se já tem tipo e campos na mesma mensagem (etapa 0)
        if self.context['step'] == 0 and ',' in user_input:
            parts = [p.strip() for p in user_input.split(',', 1)]
            type_part = parts[0]
            fields_part = parts[1] if len(parts) > 1 else ''
            
            if type_part in ['docx', 'doc']:
                self.context['template_type'] = 'docx'
            elif type_part in ['xlsx', 'xls']:
                self.context['template_type'] = 'xlsx'
            else:
                return "Tipo inválido! Use DOC/DOCX ou XLSX"
            
            if fields_part:
                self.context['step'] = 1
                return self._handle_fields_input(fields_part)
            else:
                self.context['step'] = 1
                return self.get_format_help(self.context['template_type'])
        
        # Etapas do fluxo normal
        if self.context['step'] == 0:
            return self._handle_type_selection(user_input)
        
        if self.context['step'] == 1:
            return self._handle_fields_input(user_input)
        
        if self.context['step'] == 2:
            return self._handle_confirmation(user_input)
        
        return "Comando não reconhecido. Digite 'ajuda' para ver as opções."
    
    
    def _generate_and_save_template(self) -> str:
        """Gera o template e usa o método principal"""
        template_type = self.context['template_type']
        fields = self.context['fields']

        defaultext = '.docx' if template_type == 'docx' else '.xlsx'
        filetypes = [('Documento Word', '*.docx')] if template_type == 'docx' else [('Planilha Excel', '*.xlsx')]
        default_name = f"template_{datetime.now().strftime('%Y%m%d')}{defaultext}"

        output_path = filedialog.asksaveasfilename(
            parent=self.master,
            title="Salvar Template Como",
            defaultextension=defaultext,
            filetypes=filetypes,
            initialfile=default_name
        )

        if not output_path:
            self.reset_context()
            return "❌ Operação cancelada - nenhum arquivo foi salvo"

        try:
            processor = DocumentProcessor(self.master)
            content_path = processor.generate_template(template_type, fields)
            if os.path.exists(content_path):
                os.replace(content_path, output_path)
            self.reset_context()
            return f"✅ Template gerado com sucesso!\n📂 Salvo em: {output_path}"
        except Exception as e:
            self.reset_context()
            return f"❌ Erro ao salvar arquivo: {str(e)}"
    def _handle_confirmation(self, user_input: str) -> str:
        """Processa a confirmação final"""
        if user_input.startswith(('s', 'sim', 'confirmar')):
            result = self._generate_and_save_template()
            self.reset_context()
            return result
        
        if user_input.startswith(('n', 'não', 'cancelar')):
            self.reset_context()
            return "❌ Operação cancelada. Posso ajudar com algo mais?"
        
        return "⚠️ Por favor, confirme com 'sim' ou 'não'"
    
    def _handle_fields_input(self, user_input: str) -> str:
        """Processa a entrada dos campos com formatação"""
        fields = []
        for item in user_input.split(','):
            item = item.strip()
            if item:
                # Processa formatação se existir
                if '*' in item:
                    field_name, fmt = item.split('*', 1)  # Divide apenas no primeiro *
                    fmt = fmt.strip().lower()
                    # Mapeia abreviações para formatos completos
                    fmt = self.style_shortcuts.get(fmt, fmt)
                    self.context['styles'][field_name.strip()] = fmt
                    fields.append(field_name.strip())
                else:
                    fields.append(item)
        
        if not fields:
            return "❌ Nenhum campo válido detectado. Por favor, informe os campos novamente."
        
        self.context['fields'] = fields
        self.context['step'] = 2
        return self._generate_preview()
    
    def _handle_type_selection(self, user_input: str) -> str:
        user_input = user_input.lower().strip()
        
        # Já veio com tipo e campos juntos?
        if ',' in user_input:
            parts = [p.strip() for p in user_input.split(',', 1)]  # Split apenas na primeira vírgula
            template_type = parts[0]
            fields = parts[1] if len(parts) > 1 else ''
            
            if template_type in ['docx', 'doc']:
                self.context['template_type'] = 'docx'
                if fields:
                    return self._handle_fields_input(fields)
                return self.get_format_help('docx')
            
            elif template_type in ['xlsx', 'xls']:
                self.context['template_type'] = 'xlsx'
                if fields:
                    return self._handle_fields_input(fields)
                return self.get_format_help('xlsx')
        
        # Processamento normal do tipo
        if user_input in ['docx', 'doc']:
            self.context['template_type'] = 'docx'
            self.context['step'] = 1
            return self.get_format_help('docx')
        
        if user_input in ['xlsx', 'xls']:
            self.context['template_type'] = 'xlsx'
            self.context['step'] = 1
            return self.get_format_help('xlsx')
        
        # Resposta para tipo inválido
        return (
            "⚠️ Tipo de template não reconhecido. Por favor, escolha:\n"
            "- 'docx' para Documento Word\n"
            "- 'xlsx' para Planilha Excel\n\n"
            "Digite o tipo de template que deseja criar:"
        )
    
    def _generate_preview(self) -> str:
        """Gera uma pré-visualização mais rica do template"""
        template_type = self.context['template_type']
        preview = []
        
        if template_type == 'docx':
            preview.append("📄 DOCUMENTO WORD TEMPLATE")
            preview.append("="*40)
            preview.append("\n[Cabeçalho Automático]")
            preview.append("\n# TÍTULO PRINCIPAL")
            
            for field in self.context['fields']:
                style = self.context['styles'].get(field, 'padrão')
                if style == 'heading1':
                    preview.append(f"\n# {field.upper()}")
                elif style == 'heading2':
                    preview.append(f"\n## {field.title()}")
                elif style == 'bullet':
                    preview.append(f"\n • {field}: _______________")
                elif style == 'quote':
                    preview.append(f"\n> {field}: _______________")
                else:
                    preview.append(f"\n{field}: _______________")
            
            preview.append("\n" + "="*40)
            preview.append("[Rodapé Automático]")
        
        else:  # XLSX
            preview.append("📊 PLANILHA EXCEL TEMPLATE")
            preview.append("="*40)
            
            # Cabeçalhos
            headers = []
            for field in self.context['fields']:
                style = self.context['styles'].get(field, '')
                if style:
                    headers.append(f"{field} ({style})")
                else:
                    headers.append(field)
            
            preview.append("\n| " + " | ".join(headers) + " |")
            preview.append("|" + "|".join(["---"]*len(headers)) + "|")
            
            # 3 linhas de exemplo
            for _ in range(3):
                row = []
                for field in self.context['fields']:
                    style = self.context['styles'].get(field, '')
                    if style == 'currency':
                        row.append("R$ 0,00")
                    elif style == 'date':
                        row.append(datetime.now().strftime("%d/%m/%Y"))
                    elif style == 'formula':
                        row.append("=SOMA(...)")
                    else:
                        row.append("___________")
                
                preview.append("| " + " | ".join(row) + " |")
            
            preview.append("\n" + "="*40)
            preview.append(f"Total de colunas: {len(self.context['fields'])}")
        
        return "\n".join(preview)
        
    def _format_styles_info(self) -> str:
        """Formata informações de estilos para exibição"""
        if not self.context['styles']:
            return "Nenhum estilo especial aplicado"
        
        style_lines = []
        for field, style in self.context['styles'].items():
            style_lines.append(f"• {field}: {style}")
        
        return "\n".join(style_lines)

    def reset_context(self):
        """Reinicia o contexto para o estado inicial"""
        self.context = {
            'step': 0,
            'template_type': None,
            'fields': []
        }


    def generate_preview(self) -> str:
        """Gera uma pré-visualização mais rica do template"""
        template_type = self.context['template_type']
        preview = []
        
        if template_type == 'docx':
            preview.append("📄 DOCUMENTO WORD TEMPLATE")
            preview.append("="*40)
            preview.append("\n[Cabeçalho Automático]")
            preview.append("\n# TÍTULO PRINCIPAL")
            
            for field in self.context['fields']:
                style = self.context['styles'].get(field, 'padrão')
                if style == 'heading1':
                    preview.append(f"\n# {field.upper()}")
                elif style == 'heading2':
                    preview.append(f"\n## {field.title()}")
                elif style == 'bullet':
                    preview.append(f"\n • {field}: _______________")
                elif style == 'quote':
                    preview.append(f"\n> {field}: _______________")
                else:
                    preview.append(f"\n{field}: _______________")
            
            preview.append("\n" + "="*40)
            preview.append("[Rodapé Automático]")
        
        else:  # XLSX
            preview.append("📊 PLANILHA EXCEL TEMPLATE")
            preview.append("="*40)
            
            # Cabeçalhos
            headers = []
            for field in self.context['fields']:
                style = self.context['styles'].get(field, '')
                if style:
                    headers.append(f"{field} ({style})")
                else:
                    headers.append(field)
            
            preview.append("\n| " + " | ".join(headers) + " |")
            preview.append("|" + "|".join(["---"]*len(headers)) + "|")
            
            # 3 linhas de exemplo
            for _ in range(3):
                row = []
                for field in self.context['fields']:
                    style = self.context['styles'].get(field, '')
                    if style == 'currency':
                        row.append("R$ 0,00")
                    elif style == 'date':
                        row.append(datetime.now().strftime("%d/%m/%Y"))
                    elif style == 'formula':
                        row.append("=SOMA(...)")
                    else:
                        row.append("___________")
                
                preview.append("| " + " | ".join(row) + " |")
            
            preview.append("\n" + "="*40)
            preview.append(f"Total de colunas: {len(self.context['fields'])}")
        
        return "\n".join(preview)
    
    
        
    def _validate_fields(self, fields: list) -> tuple:
        """Valida a lista de campos e retorna (status, mensagem_erro)"""
        if not fields:
            return (False, "⚠️ Nenhum campo foi informado")
        
        if len(fields) > 20:
            return (False, "⚠️ Limite de 20 campos excedido")
        if invalid_chars.search(field):
            return (False, f"Caracteres inválidos em: '{field}'")
        
        invalid_chars = re.compile(r'[\\/*?\[\]:]')
        for field in fields:
            if invalid_chars.search(field):
                return (False, f"⚠️ Nome de campo inválido: '{field}'. Não use: \\ / * ? [ ] :")
            
            if len(field) > 50:
                return (False, f"⚠️ Nome muito longo: '{field}'. Máx: 50 caracteres")
        
        return (True, "")
    
    
    def _generate_and_save_template(self) -> str:
        """Gera o template e abre diálogo para salvar arquivo"""
        template_type = self.context['template_type']
        fields = self.context['fields']
        
        filetypes = [
            ('Documento Word', '*.docx') if template_type == 'docx' 
            else ('Planilha Excel', '*.xlsx')
        ]
        defaultext = '.docx' if template_type == 'docx' else '.xlsx'
        default_name = f"template_{datetime.now().strftime('%Y%m%d')}{defaultext}"

        # Usa self.master como janela pai para o diálogo
        output_path = filedialog.asksaveasfilename(
            parent=self.master,  # Usa parent em vez de master para melhor compatibilidade
            title="Salvar Template Como",
            defaultextension=defaultext,
            filetypes=filetypes,
            initialfile=default_name
        )

        if not output_path:  # Usuário cancelou
            self.context = {'step': 0, 'template_type': None, 'fields': []}
            return "❌ Operação cancelada - nenhum arquivo foi salvo"

        # Gera o template
        try:
            if template_type == 'docx':
                doc = Document()
                doc.add_heading("Template Gerado", level=1)
                for field in fields:
                    doc.add_paragraph(f"{field}: ____________________")
                doc.save(output_path)
            else:
                wb = Workbook()
                ws = wb.active
                ws.title = "Template"
                for col_num, field in enumerate(fields, start=1):
                    ws.cell(row=1, column=col_num, value=field)
                wb.save(output_path)

            # Reseta o contexto após sucesso
            self.context = {'step': 0, 'template_type': None, 'fields': []}
            return f"✅ Template gerado com sucesso!\n📂 Salvo em: {output_path}"

        except Exception as e:
            self.context = {'step': 0, 'template_type': None, 'fields': []}
            return f"❌ Erro ao salvar arquivo: {str(e)}"
    

    

class Application:
    def __init__(self):
        
        self.root = tk.Tk()
        self.root.title("DataScan")
        self.root.geometry("1200x800")
        self.processor = DocumentProcessor(self.root)
        
        
            
        
        
        self.root.configure(bg=Config.BACKGROUND_COLOR)
        self.status_var = tk.StringVar()
        self.process_file_btn = None
        self.scrape_btn = None
        self.chat_btn = None
        self.template_btn = None
        self.chat_text = None
        self.input_entry = None
        
        self.processor = DocumentProcessor()
        self.ui_queue = Queue()
        self.progress_window = None
        self.progress_bar = None
        
        self.setup_ui()
        self.setup_ui_handler()
        self.current_chart_window = None
        self.tree = None
        self.setup_db_explorer()
        # Carregar modelo em thread separada
        self.loading_complete = False
        threading.Thread(target=self.initialize_model, daemon=True).start()
        
   
    
    def initialize_model(self):
        try:
            self.processor = DocumentProcessor()
            self.loading_complete = True
            print("Modelo carregado com sucesso!")
        except Exception as e:
            messagebox.showerror("Erro Fatal", f"Falha ao carregar IA: {str(e)}")
            self.root.quit()
            
            
            
    def process_ai_request(self, query: str):
        """Processa a solicitação do usuário com múltiplas camadas de validação."""
        try:
            # Verifica se a query está vazia
            if not query.strip():
                self.update_chat("Erro: Por favor, digite uma solicitação válida.", "error")
                return

            # Gera a resposta crua da IA
            raw_response = self.processor.generate_ai_response(query)
            
            # Camadas de filtragem e pós-processamento
            processed_response = self._filter_ai_response(raw_response, query)
            
            # Atualiza o chat com a resposta processada
            self.update_chat(f"IA: {processed_response}", "ai")

        except Exception as e:
            self.update_chat(f"Erro no processamento: {str(e)}", "error")

    def _filter_ai_response(self, response: str, original_query: str) -> str:
        """Aplica múltiplos filtros para garantir relevância na resposta."""
        # Filtro 1: Respostas muito curtas
        if len(response.split()) < 3:
            return (
                "Não entendi completamente. Poderia reformular sua solicitação incluindo:\n"
                "- Tipo de template (DOC/XLSX)\n"
                "- Campos necessários\n"
                "- Exemplo de uso pretendido"
            )

        # Filtro 2: Conteúdo fora do contexto de templates
        keywords = ['template', 'doc', 'xlsx', 'campo', 'planilha', 'documento']
        if not any(kw in response.lower() for kw in keywords):
            return self._get_guided_template_prompt(original_query)

        # Filtro 3: Respostas demasiado genéricas
        generic_phrases = ['como tal', 'através disso', 'neste contexto', 'deve-se considerar']
        if any(phrase in response.lower() for phrase in generic_phrases):
            return (
                "Por favor, especifique detalhes técnicos:\n"
                "1. Tipo de documento\n"
                "2. Nomes exatos dos campos\n"
                "3. Exemplo de estrutura desejada"
            )

        # Normalização final da resposta
        return self._normalize_response(response)
            
    
    def _get_guided_template_prompt(self, query: str) -> str:
        """Retorna um prompt estruturado quando detecta falta de informações."""
        # Usando uma string multilinha com aspas triplas
        base_prompt = """Para criar um template, preciso das seguintes informações:

    ➤ Tipo: 
    - DOC (Documento Word)
    - XLSX (Planilha Excel)

    ➤ Campos: 
    Lista separada por vírgulas
    Ex: Nome, Data, Valor

    ➤ Formatação: 
    (Opcional) Requisitos específicos de layout

    Por favor, forneça esses detalhes."""
        
        return base_prompt

    def _normalize_response(self, response: str) -> str:
        """Normaliza a resposta técnica para manter consistência."""
        replacements = {
            'documento word': 'DOC',
            'planilha excel': 'XLSX',
            'campos': '→ Campos',
            'modelo': 'Template'
        }
        
        for term, replacement in replacements.items():
            # CORRIJA ESTA LINHA ↓
            response = re.sub(rf'\b{re.escape(term)}\b', replacement, response, flags=re.IGNORECASE)
        
        # Remove markdown não intencional
        response = re.sub(r'([*_`])', '', response)
        
        return response.strip()
    
    def _validate_response_context(self, response: str) -> bool:
        """Garante que a resposta está relacionada a templates"""
        required_terms = [
            'template', 'doc', 'xlsx', 'campo', 
            'documento', 'planilha', 'formulário'
        ]
        return any(term in response.lower() for term in required_terms)
    



    def setup_ui(self):
        main_frame = tk.Frame(self.root, bg=Config.BACKGROUND_COLOR)
        main_frame.pack(fill=tk.BOTH, expand=True, padx=20, pady=20)

        # Frame centralizado
        center_frame = tk.Frame(main_frame, bg=Config.BACKGROUND_COLOR)
        center_frame.place(relx=0.5, rely=0.5, anchor='center')

        # Título moderno
        title_frame = tk.Frame(center_frame, bg=Config.BACKGROUND_COLOR)
        title_frame.pack(pady=20)
        
        self.title_label = tk.Label(
            title_frame,
            text="DataScan",
            font=("Segoe UI", 34, "bold"),
            fg=Config.PRIMARY_COLOR,
            bg=Config.BACKGROUND_COLOR
        )
        self.title_label.pack()
        
        self.subtext_label = tk.Label(
            title_frame,
            text="Intelligent Data Processing",
            font=("Segoe UI", 18),
            fg=Config.SECONDARY_COLOR,
            bg=Config.BACKGROUND_COLOR
        )
        self.subtext_label.pack()
        
            # Adicionar barra de status
        self.status_bar = tk.Label(
            self.root, 
            textvariable=self.status_var,
            bg=Config.SECONDARY_COLOR,
            fg=Config.TEXT_COLOR,
            anchor=tk.W
        )
        self.status_bar.pack(side=tk.BOTTOM, fill=tk.X)
        
        self.subtext2_label = tk.Label(
            title_frame,
            text="By: github.com/Marlon009",
            font=("Segoe UI", 12, "italic"),
            fg=Config.SECONDARY_COLOR,
            bg=Config.BACKGROUND_COLOR
        )
        self.subtext2_label.pack()

        # Botões principais
        button_frame = tk.Frame(center_frame, bg=Config.BACKGROUND_COLOR)
        button_frame.pack(pady=30)

        button_style = {
            'font': ('Segoe UI', 12),
            'width': 20,
            'height': 2,
            'bg': Config.PRIMARY_COLOR,
            'fg': Config.TEXT_COLOR,
            'activebackground': Config.HOVER_COLOR,
            'activeforeground': Config.TEXT_COLOR,
            'relief': 'flat',
            'borderwidth': 0,
            'cursor': 'hand2'
        }

        buttons = [
            ("📁 Processar Arquivo", self.process_file),
            ("🌐 Scraping Web", self.scrape_website),
            ("💬 Terminal IA", self.chat_with_ai),
            ("💾 Backup", self.create_backup)
        ]

        for text, cmd in buttons:
            btn = tk.Button(button_frame, text=text, compound=tk.LEFT, **button_style)
            btn.config(command=cmd)
            btn.bind("<Enter>", lambda e: e.widget.config(bg=Config.ACCENT_COLOR))
            btn.bind("<Leave>", lambda e: e.widget.config(bg=Config.PRIMARY_COLOR))
            btn.pack(side=tk.LEFT, padx=10, pady=5)
            
            
    def is_command(self, query: str) -> bool:
        """Verifica se a entrada do usuário é um comando especial."""
        command_map = {
            'criar template': self.handle_template_creation,
            'processar arquivo': self.process_file,
            'analisar site': self.scrape_website,
            'ajuda': self.show_help,
            'limpar': self.clear_chat_context,
            'sair': lambda: self.root.quit()
        }
        
        return any(cmd in query.lower() for cmd in command_map.keys())
    
    # Na classe Application, adicione este método:
    def clean_input(self, text: str) -> str:
        """Limpa e padroniza o input do usuário"""
        # Remove espaços extras e caracteres especiais
        cleaned = text.strip()
        # Remove múltiplos espaços internos
        cleaned = ' '.join(cleaned.split())
        return cleaned      

    def show_progress(self, message: str):
        if self.progress_window is None:
            self.progress_window = tk.Toplevel(self.root)
            self.progress_window.title("Processando")
            self.progress_window.configure(bg='#0a0a0a')
            
            self.progress_label = tk.Label(
                self.progress_window,
                text=message,
                font=("Segoe UI", 12),
                fg="#00ff00",
                bg='#0a0a0a'
            )
            self.progress_label.pack(padx=30, pady=15)
            
            self.progress_bar = ttk.Progressbar(
                self.progress_window,
                orient=tk.HORIZONTAL,
                length=400,
                mode='indeterminate',
                style='green.Horizontal.TProgressbar'
            )
            self.progress_bar.pack(pady=15)
            
            style = ttk.Style()
            style.theme_use('clam')
            style.configure(
                'green.Horizontal.TProgressbar',
                background='#00ff00',
                troughcolor='#002200'
            )
            
            self.progress_window.grab_set()

    def update_progress(self, value: int, message: str = None):
        """Atualiza o valor da barra de progresso e a mensagem."""
        if self.progress_bar:
            self.progress_bar['value'] = value
            if message and self.progress_label:
                self.progress_label.config(text=message)
            self.progress_window.update_idletasks()

    def hide_progress(self):
        """Fecha a janela de progresso."""
        if self.progress_window:
            self.progress_window.grab_release()
            self.progress_window.destroy()
            self.progress_window = None
            self.progress_bar = None
            self.progress_label = None

    def process_file(self):
        """Inicia o processamento de arquivo com barra de progresso."""
        self.ui_command('disable_buttons')
        self.ui_command('update_status', "Aguardando seleção de arquivo...")
        file_path = filedialog.askopenfilename(
            filetypes=[("Documentos", "*.pdf *.docx *.xlsx *.txt"), ("Todos", "*.*")]
        )
        if file_path:
            self.ui_command('show_progress', "Processando arquivo...")
            threading.Thread(target=self._process_file, args=(file_path,), daemon=True).start()
        else:
            self.ui_command('enable_buttons')
            self.ui_command('update_status', "Operação cancelada")

    def _process_file(self, file_path: str):
        """Simula o processamento de um arquivo com barra de progresso."""
        try:
            total_steps = 100
            for i in range(total_steps + 1):
                time.sleep(0.05)  # Simula um processamento demorado
                progress = int((i / total_steps) * 100)
                self.ui_command('update_progress', progress, f"Processando... {progress}%")
            
            content = self.processor.extract_from_file(file_path)
            self.processor.document_content = content

            self.ui_command('update_status', "Selecione o template...")
            template_path = filedialog.askopenfilename(
                filetypes=[("Templates", "*.docx *.xlsx *.pdf")]
            )
            if template_path:
                self.ui_command('update_status', "Gerando saída...")
                output_path = filedialog.asksaveasfilename(
                    defaultextension=os.path.splitext(template_path)[1]
                )
                if output_path:
                    self.processor.fill_template(template_path, output_path, {"content": content})
                    self.ui_command('show_info', "Arquivo processado com sucesso!")
        except Exception as e:
            self.ui_command('show_error', str(e))
        finally:
            self.ui_command('hide_progress')
            self.ui_command('update_status', "Pronto para novas operações")
            self.ui_command('enable_buttons')

    def setup_ui_handler(self):
        def check_queue():
            try:
                while not self.ui_queue.empty():
                    action, *args = self.ui_queue.get_nowait()
                    handler = {
                        'show_progress': self.show_progress,
                        'update_progress': self.update_progress,
                        'hide_progress': self.hide_progress,
                        'show_error': self.show_error,
                        'show_info': self.show_info,
                        'update_status': lambda x: self.status_var.set(f"» {x}"),
                        'enable_buttons': self.enable_buttons,
                        'disable_buttons': self.disable_buttons
                    }.get(action)
                    
                    if handler:
                        handler(*args) if args else handler()
                        
            except Exception as e:
                print(f"Erro na fila: {str(e)}")
            finally:
                self.root.after(100, check_queue)

        self.root.after(100, check_queue)


    def ui_command(self, action, *args):
        self.ui_queue.put((action, *args))
            
    def hide_progress(self):
        if self.progress_window:
            self.progress_bar.stop()
            self.progress_window.grab_release()
            self.progress_window.destroy()
            self.progress_window = None
            self.progress_bar = None

    def show_error(self, message):
        messagebox.showerror("Erro de Sistema", message)
        self.ui_command('update_status', f"Erro: {message}")

    def show_info(self, message):
        messagebox.showinfo("Operação Concluída", message)
        self.ui_command('update_status', message)

    def enable_buttons(self):
        if self.process_file_btn and self.scrape_btn and self.chat_btn and self.template_btn:
            self.process_file_btn.config(state=tk.NORMAL)
            self.scrape_btn.config(state=tk.NORMAL)
            self.chat_btn.config(state=tk.NORMAL)
            self.template_btn.config(state=tk.NORMAL)

    def disable_buttons(self):
        if self.process_file_btn and self.scrape_btn and self.chat_btn and self.template_btn:
            self.process_file_btn.config(state=tk.DISABLED)
            self.scrape_btn.config(state=tk.DISABLED)
            self.chat_btn.config(state=tk.DISABLED)
            self.template_btn.config(state=tk.DISABLED)

    def process_file(self):
        self.ui_command('disable_buttons')
        self.ui_command('update_status', "Aguardando seleção de arquivo...")
        file_path = filedialog.askopenfilename(
            filetypes=[("Documentos", "*.pdf *.docx *.xlsx *.txt"), ("Todos", "*.*")]
        )
        if file_path:
            self.ui_command('update_status', "Processando arquivo...")
            threading.Thread(target=self._process_file, args=(file_path,), daemon=True).start()
        else:
            self.ui_command('enable_buttons')
            self.ui_command('update_status', "Operação cancelada")

    def _process_file(self, file_path: str):
        try:
            self.ui_command('show_progress', "Decodificando arquivo...")
            content = self.processor.extract_from_file(file_path)
            self.processor.document_content = content

            data = {"content": content}

            self.ui_command('update_status', "Selecione o template...")
            template_path = filedialog.askopenfilename(
                filetypes=[("Templates", "*.docx *.xlsx *.pdf")]
            )
            if template_path:
                self.ui_command('update_status', "Gerando saída...")
                output_path = filedialog.asksaveasfilename(
                    defaultextension=os.path.splitext(template_path)[1]
                )
                if output_path:
                    self.processor.fill_template(template_path, output_path, data)
                    self.ui_command('show_info', "Arquivo processado com sucesso!")
        except Exception as e:
            self.ui_command('show_error', str(e))
        finally:
            self.ui_command('hide_progress')
            self.ui_command('update_status', "Pronto para novas operações")
            self.ui_command('enable_buttons')


    def scrape_website(self):
        try:
            self.ui_command('disable_buttons')
            self.ui_command('update_status', "Aguardando URL...")

            url = simpledialog.askstring("Scraping Web", "Digite a URL do site:")
            if not url:
                self.ui_command('enable_buttons')
                self.ui_command('update_status', "Operação cancelada")
                return

            # Adicionar validação completa
            if not url.startswith(('http://', 'https://')):
                url = f'https://{url}'

            if not validators.url(url):
                raise ValueError("URL inválida")

            self.ui_command('update_status', "Analisando site...")

            # Adicionar headers para evitar bloqueio
            headers = {
                'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36'
            }

            response = requests.get(url, headers=headers, timeout=15)
            response.raise_for_status()

            # Executa o scraping em uma thread separada
            threading.Thread(target=self._scrape_and_process, args=(url,), daemon=True).start()

        except requests.exceptions.RequestException as e:
            self.ui_command('update_status', f"Erro ao acessar o site: {e}")
            self.ui_command('enable_buttons')
        except ValueError as e:
            self.ui_command('update_status', str(e))
            self.ui_command('enable_buttons')


    def _scrape_and_process(self, url: str):
        try:
            self.ui_command('show_progress', "Coletando dados do site...")
            data = self.processor.scrape_website(url)
            self.processor.scraped_data = data

            self.ui_command('update_status', "Selecione o template...")
            template_path = filedialog.askopenfilename(
                filetypes=[("Templates", "*.docx *.xlsx *.pdf")]
            )
            if template_path:
                self.ui_command('update_status', "Gerando relatório...")
                output_path = filedialog.asksaveasfilename(
                    defaultextension=os.path.splitext(template_path)[1]
                )
                if output_path:
                    self.processor.fill_template(template_path, output_path, data)
                    self.ui_command('show_info', "Dados processados com sucesso!")
        except Exception as e:
            self.ui_queue.put(('show_error', str(e)))
        finally:
            self.ui_queue.put(('hide_progress',))
            self.ui_queue.put(('update_status', "Pronto para novas operações"))
            self.ui_queue.put(('enable_buttons',))

    

    def _generate_template_thread(self, template_type: str, fields: list):
        try:
            self.ui_command('show_progress', "Gerando template com IA...")
            output_path = self.processor.generate_template(template_type, fields)
            self.ui_command('show_info', f"Template criado: {output_path}")
        except Exception as e:
            self.ui_command('show_error', str(e))
        finally:
            self.ui_command('hide_progress')
            self.ui_command('enable_buttons')

    def chat_with_ai(self):
        chat_window = tk.Toplevel(self.root)
        chat_window.title("Assistente IA")
        chat_window.geometry("800x600")
        chat_window.configure(bg=Config.CHAT_BG)

        main_frame = tk.Frame(chat_window, bg=Config.CHAT_BG)
        main_frame.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)

        self.chat_text = tk.Text(
            main_frame,
            wrap=tk.WORD,
            bg=Config.CHAT_BG,
            fg=Config.TEXT_COLOR,
            font=('Segoe UI', 12),
            state=tk.DISABLED
        )
        self.chat_text.pack(fill=tk.BOTH, expand=True)

        input_frame = tk.Frame(main_frame, bg=Config.CHAT_BG)
        input_frame.pack(fill=tk.X, pady=10)

        self.input_entry = tk.Entry(
            input_frame,
            bg=Config.SECONDARY_COLOR,
            fg=Config.TEXT_COLOR,
            font=('Segoe UI', 12),
            width=50
        )
        self.input_entry.pack(side=tk.LEFT, fill=tk.X, expand=True, padx=5)

        # Botão de enviar
        send_btn = tk.Button(
            input_frame,
            text="➤",
            command=lambda: self.send_message(),
            bg=Config.ACCENT_COLOR,
            fg=Config.TEXT_COLOR,
            font=('Segoe UI', 12)
        )
        send_btn.pack(side=tk.RIGHT, padx=5)

        # Botão de limpar
        clear_btn = tk.Button(
            input_frame,
            text="🧹",
            command=self.clear_chat_context,
            bg=Config.ACCENT_COLOR,
            fg=Config.TEXT_COLOR,
            font=('Segoe UI', 12)
        )
        clear_btn.pack(side=tk.RIGHT, padx=5)

        self.input_entry.bind("<Return>", lambda e: self.send_message())

        # Mensagem inicial - Corrigido com msg_type
        self.update_chat(
            "╭──────────────────────────────────────────╮\n"
            "│          ASSISTENTE DATASCAN 1.0         │\n"
            "╰──────────────────────────────────────────╯\n\n"
            "📌 **Como usar**:\n"
            "1. Digite 'docx' ou 'xlsx' para escolher o tipo\n"
            "2. Liste os campos separados por vírgulas\n"
            "3. Adicione formatações com * (opcional)\n\n"
            "💡 **Exemplo**: 'Título*h1, Itens*blt, Data'\n"
            "────────────────────────────────────────────\n"
            "Digite 'ajuda' a qualquer momento para orientações",
            "system"  # Tipo de mensagem adicionado
        )
        
        self.input_entry.bind("<Return>", lambda e: self.send_message())
        
        
    def send_message(self, _event=None):  # Adicione este método à classe
        user_input = self.input_entry.get().strip()
        if user_input:
            try:
                self.update_chat(f"Você: {user_input}", "user")
                threading.Thread(
                    target=self.process_ai_request,
                    args=(user_input,),
                    daemon=True
                ).start()
            except Exception as e:
                self.update_chat(f"Erro: {str(e)}", "error")
            finally:
                self.input_entry.delete(0, tk.END)
        
    def upload_file_for_chat(self, parent_window):
        file_path = filedialog.askopenfilename(
            filetypes=[("Documentos", "*.pdf *.docx *.xlsx *.txt"), ("Todos", "*.*")]
        )
        if file_path:
            try:
                content = self.processor.extract_from_file(file_path)
                self.processor.document_content = content
                self.update_chat(f"Sistema: Arquivo carregado com sucesso: {os.path.basename(file_path)}\n", "system")
                self.update_chat(f"Resumo do documento:\n{content[:500]}...\n\n", "file")
            except Exception as e:
                self.update_chat(f"Erro ao processar arquivo: {str(e)}\n", "error")

    def clear_chat_context(self):
        self.processor.chat_context = []  # Limpa o contexto
        self.update_chat("Sistema: Contexto da conversa limpo com sucesso\n", "system")

    def process_query(self, window):
        query = self.input_entry.get().strip()
        self.input_entry.delete(0, tk.END)
        if not query:
            return

        # Aplicar limpeza no input
        query = self.clean_input(query)

        self.update_chat(f"Você: {query}\n", "user")

        try:
            if self.is_command(query):
                self.handle_command(query)
            else:
                threading.Thread(
                    target=self.generate_context_aware_response,
                    args=(query,),
                    daemon=True
                ).start()

        except Exception as e:
            self.update_chat(f"Erro: {str(e)}\n", "error")
                
    def generate_context_aware_response(self, query: str):
        """Gera uma resposta contextualizada com formatação correta"""
        try:
            start_time = time.time()
            
            # Gera resposta crua da IA
            raw_response = self.processor.generate_ai_response(query)
            
            # Processa a resposta para remover formatação indesejada
            cleaned_response = self.clean_ai_output(raw_response)
            
            elapsed_time = time.time() - start_time
            
            # Formatação profissional da resposta
            formatted_response = (
                f"\n[IA • {elapsed_time:.2f}s]\n"
                f"➤ {cleaned_response}\n"
                f"{'─' * 60}\n"
            )
            
            self.update_chat(formatted_response, "ai")
            
        except Exception as e:
            self.update_chat(f"\n❌ Erro: {str(e)}\n", "error")
            
            
    def clean_ai_output(self, text: str) -> str:
        """Remove artefatos de formatação indesejados"""
        patterns = [
            r'<.*?>',          # Tags HTML
            r'\b\d+\.\s*',     # Listas numeradas
            r'[-=]{4,}',       # Linhas divisórias
            r'Resposta:\s*',   # Marcadores de resposta
            r'^[\W_]+|[\W_]+$' # Caracteres especiais no início/fim
        ]
        
        for pattern in patterns:
            text = re.sub(pattern, '', text, flags=re.IGNORECASE|re.MULTILINE)
        
        # Correção de formatação
        text = text.replace('  ', ' ').strip()
        
        if len(text) > 0:
            # Capitalização inteligente
            text = text[0].upper() + text[1:]
            if text[-1] not in {'.', '!', '?'}:
                text += '.'
                
        return text
            
    def handle_command(self, query: str):
        """Executa comandos especiais."""
        command_map = {
            '/criar template': self.handle_template_creation,
            '/processar arquivo': self.process_file,  # Alterado para process_file
            '/analisar site': self.scrape_website,    # Alterado para scrape_website
            '/ajuda': self.show_help,
            '/limpar': self.clear_chat_context,
            '/sair': lambda: self.root.quit()
        }
        
        for cmd, handler in command_map.items():
            if cmd in query.lower():
                handler()
                return


    def handle_template_creation(self, query: str):
        try:
            self.update_chat("Assistente: Analisando sua solicitação...\n", "system")
            
            prompt = (
                f"Com base na seguinte solicitação: '{query}'\n"
                "Identifique:\n"
                "1. Tipo de documento (docx/xlsx)\n"
                "2. Campos necessários\n"
                "3. Formatação especial\n"
                "Retorne no formato JSON:"
                '{"type": "docx", "fields": ["nome", "data"], "instructions": "..."}'
            )
            
            response = self.processor.generate_ai_response(prompt)
            params = json.loads(response)
            
            self.update_chat(
                f"Assistente: Criando template {params['type'].upper()} com campos: "
                f"{', '.join(params['fields'])}\n",
                "system"
            )
            
            template_path = self.processor.generate_template(
                params['type'],
                params['fields']
            )
            
            self.update_chat(
                f"✅ Template criado com sucesso!\n"
                f"📂 Caminho: {template_path}\n"
                f"📝 Instruções: {params.get('instructions', 'Nenhuma')}\n\n",
                "success"
            )
            
        except Exception as e:
            self.update_chat(f"❌ Erro na criação: {str(e)}\n", "error")


    def generate_response(self, query, window):
        try:
            start_time = time.time()
            response = self.processor.generate_ai_response(query)
            elapsed_time = time.time() - start_time
            
            # Formatar resposta corretamente
            formatted_response = f"IA ({elapsed_time:.2f}s):\n{response}\n"
            formatted_response += "―" * 50 + "\n"
            
            self.update_chat(formatted_response, "ai")
        except Exception as e:
            self.update_chat(f"Erro: {str(e)}\n", "error")
            
            
    def show_help(self):
        help_text = """Comandos disponíveis:
/ajuda - Mostra esta mensagem
/limpar - Limpa o contexto da conversa
/sair - Fecha o chat
/carregar - Abre diálogo para carregar arquivo
/criar template - Cria um novo template com IA
"""
        self.update_chat(f"Sistema:\n{help_text}\n", "system")

    def update_chat(self, message: str, msg_type: str = "ai") -> None:
        """Atualiza o chat com uma nova mensagem
        
        Args:
            message: Texto da mensagem
            msg_type: Tipo da mensagem ('user', 'ai', 'system', 'error')
        """
        # Mapeamento de cores para cada tipo de mensagem
        color_map = {
            "user": "#CFD8DC",      # Cinza claro
            "ai": "#C5E1A5",        # Verde claro
            "system": "#81D4FA",    # Azul claro
            "error": "#EF9A9A",     # Vermelho claro
            "success": "#A5D6A7",   # Verde suave
            "file": "#90CAF9"       # Azul pastel
        }

        # Configura o estado do chat_text para edição
        self.chat_text.config(state=tk.NORMAL)
        
        # Insere a mensagem com formatação
        timestamp = datetime.now().strftime("[%H:%M:%S] ")
        self.chat_text.insert(tk.END, timestamp, "timestamp")
        self.chat_text.insert(tk.END, f"{message}\n", msg_type)
        
        # Rola para o final do texto
        self.chat_text.see(tk.END)
        self.chat_text.config(state=tk.DISABLED)

    def setup_db_explorer(self):
        """Janela para explorar a estrutura do banco de dados"""
        self.explorer_window = tk.Toplevel(self.root)
        self.explorer_window.title("Explorador de Banco de Dados")
        self.explorer_window.geometry("600x400")
        
        self.tree = ttk.Treeview(self.explorer_window)
        self.tree.pack(fill=tk.BOTH, expand=True)
        
        vsb = ttk.Scrollbar(self.explorer_window, orient="vertical", command=self.tree.yview)
        vsb.pack(side=tk.RIGHT, fill=tk.Y)
        self.tree.configure(yscrollcommand=vsb.set)
        
        self.explorer_window.withdraw()

    def show_db_explorer(self):
        """Atualiza e exibe o explorador de banco de dados"""
        try:
            schema = self.processor.db_manager.get_schema()
            self.tree.delete(*self.tree.get_children())
            
            for table in schema['tables']:
                table_id = self.tree.insert("", "end", text=table['name'], values=["Table"])
                for column in table['columns']:
                    self.tree.insert(table_id, "end", 
                                   text=f"{column['name']} ({column['type']})",
                                   values=["Column"])
            
            self.explorer_window.deiconify()
        except Exception as e:
            self.show_error(str(e))

    def create_backup(self):
        """Cria backup compactado do diretório de processamento"""
        try:
            if not os.path.exists(Config.TEMP_DIR):
                raise ValueError("Nenhum dado para backup encontrado")
                
            backup_name = f"backup_{datetime.now().strftime('%Y%m%d_%H%M')}.zip"
            backup_path = os.path.join(Config.BACKUP_DIR, backup_name)
            
            with zipfile.ZipFile(backup_path, 'w', zipfile.ZIP_DEFLATED) as zipf:
                for root, _, files in os.walk(Config.TEMP_DIR):
                    for file in files:
                        file_path = os.path.join(root, file)
                        arcname = os.path.relpath(file_path, Config.TEMP_DIR)
                        zipf.write(file_path, arcname)
                        
            self.show_info(f"Backup criado com sucesso em:\n{backup_path}")
            return backup_path
            
        except Exception as e:
            self.show_error(f"Falha no backup: {str(e)}")
            return None

    def show_chart(self, data: dict):
        """Exibe os dados em formato gráfico"""
        if self.current_chart_window:
            self.current_chart_window.destroy()
            
        self.current_chart_window = tk.Toplevel(self.root)
        self.current_chart_window.title("Visualização de Dados")
        
        fig = plt.Figure(figsize=(6, 4), dpi=100)
        ax = fig.add_subplot(111)
        
        if isinstance(data, list) and len(data) > 0:
            df = pd.DataFrame(data)
            df.plot(kind='bar', ax=ax)
            
            canvas = FigureCanvasTkAgg(fig, master=self.current_chart_window)
            canvas.draw()
            canvas.get_tk_widget().pack(fill=tk.BOTH, expand=True)


if __name__ == "__main__":
    os.makedirs(Config.TEMP_DIR, exist_ok=True)
    os.makedirs(Config.BACKUP_DIR, exist_ok=True)
    app = Application()
    app.root.mainloop()