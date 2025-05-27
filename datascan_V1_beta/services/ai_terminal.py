import google.generativeai as genai
from docx import Document
from PyPDF2 import PdfWriter
import pandas as pd
import os
import re
from dotenv import load_dotenv

load_dotenv()

class AITerminal:
    def __init__(self, api_key=None):
        self.api_key = api_key
        self.model = None
        self.initialize_model()
        self.conversation_state = {
            'awaiting_response': False,
            'document_type': None,
            'fields': [],
            'content': None,
            'formatting': {}
        }
        self.initialize_model()
        
    def initialize_model(self):
        if self.api_key:
            try:
                genai.configure(api_key=self.api_key)
                self.model = genai.GenerativeModel('gemini-pro')
                return True
            except Exception as e:
                print(f"Erro na inicialização: {e}")
                return False
        return False

    def generate_response(self, prompt):
        if not self.model:
            return "Erro: Configure sua API Key primeiro!", None
            
        try:
            if not self.conversation_state['awaiting_response']:
                self._reset_conversation()
                return self._init_document_creation(prompt)
            else:
                return self._handle_document_details(prompt)
                
        except Exception as e:
            return f"Erro na geração: {str(e)}", None

    def _init_document_creation(self, prompt):
        doc_type = self._detect_document_type(prompt)
        if doc_type:
            self.conversation_state.update({
                'awaiting_response': True,
                'document_type': doc_type
            })
            return (
                f"Vou ajudar com o {doc_type.upper()}. Por favor, informe:\n"
                "1. Quais campos devem ser incluídos (separados por vírgula)\n"
                "2. Algum conteúdo base específico?\n"
                "3. Formatação especial (tabelas, cores, etc.)\n"
                "Exemplo: 'Nome, Data, Hora | Texto base: Contrato | Tabela com 3 colunas'",
                'awaiting_details'
            )
        else:
            response = self.model.generate_content(prompt)
            return response.text, None

    def _handle_document_details(self, user_input):
        # Processar a entrada do usuário
        details = [part.strip() for part in user_input.split('|')]
        
        # Obter campos
        if len(details) > 0:
            self.conversation_state['fields'] = [f.strip() for f in details[0].split(',')]
            
        # Obter conteúdo base
        if len(details) > 1 and ':' in details[1]:
            self.conversation_state['content'] = details[1].split(':')[-1].strip()
            
        # Obter formatação
        if len(details) > 2:
            self._process_formatting(details[2])
            
        # Gerar documento
        file_path = self._create_document()
        self._reset_conversation()
        
        return f"Documento gerado com sucesso!\nLocal: {file_path}", file_path

    def _create_document(self):
        # Implementação detalhada da criação do documento com base no state
        doc_type = self.conversation_state['document_type']
        
        if doc_type == 'docx':
            return self._create_custom_docx()
        elif doc_type == 'pdf':
            return self._create_custom_pdf()
        elif doc_type == 'excel':
            return self._create_custom_excel()

    def _create_custom_docx(self):
        doc = Document()
        
        # Cabeçalho
        if self.conversation_state['content']:
            doc.add_heading(self.conversation_state['content'], level=0)
            
        # Campos
        for field in self.conversation_state['fields']:
            doc.add_paragraph(f"{field}: ___________________")
            
        # Formatação adicional
        if 'table' in self.conversation_state['formatting']:
            cols = self.conversation_state['formatting']['table'].get('columns', 2)
            table = doc.add_table(rows=1, cols=cols)
            
        filename = "documento_personalizado.docx"
        doc.save(filename)
        return os.path.abspath(filename)

    def _reset_conversation(self):
        self.conversation_state = {
            'awaiting_response': False,
            'document_type': None,
            'fields': [],
            'content': None,
            'formatting': {}
        }

    def _detect_document_type(self, text):
        text = text.lower()
        if 'docx' in text: return 'docx'
        if 'pdf' in text: return 'pdf'
        if 'excel' in text or 'planilha' in text: return 'excel'
        return None